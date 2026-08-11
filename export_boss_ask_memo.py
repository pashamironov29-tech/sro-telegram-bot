#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def main() -> None:
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)

    def p(text: str, *, bold: bool = False) -> None:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)
        run.bold = bold

    p("К разговору с руководством: бот СРО (GOLD)", bold=True)
    p(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
    p("")
    p("Что хочу по итогам показа бота", bold=True)
    p("1. Обслуживание бота — 10 000 руб./мес.")
    p(
        "   Покрывает: Cursor (ИИ при доработках) + VPS (сервер в Москве) "
        "+ OpenRouter (ИИ в боте)."
    )
    p(
        "   Это не «зарплата за воздух», а рабочие расходы, "
        "без которых бот деградирует."
    )
    p("")
    p("2. Прибавка к зарплате — +15 000 руб.")
    p(
        "   За сопровождение сервиса: правки, сбои, доработки по просьбе, "
        "актуальность реестра/функций."
    )
    p("")
    p("3. Премия за создание бота — 50 000 руб. (разово)")
    p(
        "   За готовый рабочий продукт: меню, реестр, бланки, ИИ, VPS, "
        "документы по законности/ПДн."
    )
    p("")
    p("4. Иногда работать из дома")
    p(
        "   По согласованию: дни, когда удобнее сопровождать бот/доработки "
        "удалённо."
    )
    p("")
    p("Как коротко сказать вслух", bold=True)
    p(
        "«Бот готов и крутится в Москве. Прошу: разовую премию 50 тысяч "
        "за разработку; плюс 15 тысяч к окладу за сопровождение; и 10 тысяч "
        "в месяц на Cursor, сервер и OpenRouter — иначе инструмент не удержать. "
        "Плюс по возможности иногда из дома.»"
    )
    p("")
    p("Заметки для себя (не обязательно озвучивать)", bold=True)
    p(
        "— 10к на расходы: нормально и даже скромно "
        "(VPS+ИИ+Cursor легко съедают часть этой суммы)."
    )
    p(
        "— +15к к ЗП: зависит от текущего оклада; звучит как доплата "
        "за новую функцию, не как ультиматум."
    )
    p(
        "— 50к премия: для заказной разработки бота такого уровня это скромно; "
        "как внутренняя премия — адекватно."
    )
    p("— Из дома: мягкая просьба, лучше не увязывать жёстко с деньгами.")
    p(
        "— Не называть всё пакетом «или не отдаю бота». "
        "Сначала демо и польза, потом условия."
    )

    out = Path(r"C:\Users\User\Desktop") / "К_разговору_бот_СРО_хотелки.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
