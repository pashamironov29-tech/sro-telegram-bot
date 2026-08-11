"""Сборка OTCHET_BOT_SRO.docx из markdown-отчёта. Запуск: py build_otchet_bot_sro_docx.py"""
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor, Emu
from docx.oxml.ns import qn
import re

SRC = Path(__file__).parent / "OTCHET_BOT_SRO.md"
DST = Path(__file__).parent / "OTCHET_BOT_SRO.docx"

FONT = "Calibri"
FONT_SIZE = Pt(11)
HEADING_COLOR = RGBColor(0x1A, 0x47, 0x8A)


def set_run_font(run, size=FONT_SIZE, bold=False, color=None):
    run.font.name = FONT
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def add_heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    size = {1: Pt(16), 2: Pt(14), 3: Pt(12)}.get(level, Pt(12))
    set_run_font(run, size=size, bold=True, color=HEADING_COLOR)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    parts = re.split(r"\*\*(.+?)\*\*", text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        set_run_font(run, bold=(bold or i % 2 == 1))
        if italic:
            run.font.italic = True
    return p


def add_table(doc, header_row, data_rows):
    cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, val in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(val)
        set_run_font(run, bold=True)
        shading = cell._element.get_or_add_tcPr()
        bg = shading.makeelement(qn("w:shd"), {
            qn("w:fill"): "1A478A", qn("w:val"): "clear",
        })
        shading.append(bg)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r_idx, row_data in enumerate(data_rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            parts = re.split(r"\*\*(.+?)\*\*", val)
            for i, part in enumerate(parts):
                if not part:
                    continue
                run = cell.paragraphs[0].add_run(part)
                set_run_font(run, bold=(i % 2 == 1))

    return table


def parse_md_table(lines):
    header = [c.strip() for c in lines[0].strip("|").split("|")]
    rows = []
    for line in lines[2:]:
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    return header, rows


def build():
    md = SRC.read_text(encoding="utf-8")
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = FONT_SIZE
    style.paragraph_format.space_after = Pt(4)

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1.5)

    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), level=1)
            i += 1
            continue

        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), level=3)
            i += 1
            continue

        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), level=2)
            i += 1
            continue

        if line.startswith("|") and i + 2 < len(lines) and lines[i + 1].startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            header, rows = parse_md_table(table_lines)
            add_table(doc, header, rows)
            doc.add_paragraph()
            continue

        if line.startswith("---"):
            i += 1
            continue

        if line.startswith("- "):
            txt = line[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            parts = re.split(r"\*\*(.+?)\*\*", txt)
            for j, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                set_run_font(run, bold=(j % 2 == 1))
            i += 1
            continue

        if re.match(r"^\d+\.\s", line):
            txt = re.sub(r"^\d+\.\s", "", line).strip()
            p = doc.add_paragraph(style="List Number")
            parts = re.split(r"\*\*(.+?)\*\*", txt)
            for j, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                set_run_font(run, bold=(j % 2 == 1))
            i += 1
            continue

        if line.strip() == "":
            i += 1
            continue

        is_italic = line.startswith("*") and line.endswith("*") and not line.startswith("**")
        txt = line.strip("*").strip() if is_italic else line.strip()
        add_para(doc, txt, italic=is_italic)
        i += 1

    doc.save(str(DST))
    print(f"OK → {DST}")


if __name__ == "__main__":
    build()
