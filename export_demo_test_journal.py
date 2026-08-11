#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Журнал проверки бота перед показом — создать / дописать запись."""
from __future__ import annotations

import argparse
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT = Path(r"C:\Users\User\Desktop") / "Журнал_проверки_бота_перед_показом.docx"


def _font(run, *, bold: bool = False, size: int = 12) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def _add_p(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    _font(r, bold=bold)


def create_fresh() -> Path:
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2)
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    st.font.size = Pt(12)

    _add_p(doc, "Журнал проверки бота СРО перед показом руководству", bold=True)
    _add_p(doc, f"Дата начала: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    _add_p(doc, "Сервер: Москва 201.24.125.236")
    _add_p(doc, "")
    _add_p(
        doc,
        "Ниже — что проверяли (организации / люди / функции). "
        "Пополняется по ходу теста.",
    )
    _add_p(doc, "")
    _add_p(doc, "--- Записи ---", bold=True)
    doc.save(OUT)
    return OUT


def append_entry(text: str) -> Path:
    if not OUT.exists():
        create_fresh()
    doc = Document(str(OUT))
    stamp = datetime.now().strftime("%H:%M")
    _add_p(doc, "")
    _add_p(doc, f"[{stamp}] {text.strip()}")
    doc.save(OUT)
    return OUT


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--init", action="store_true")
    ap.add_argument("--add", default="")
    args = ap.parse_args()
    if args.init or not OUT.exists():
        create_fresh()
        print("created", OUT)
    if args.add.strip():
        append_entry(args.add)
        print("appended", OUT)


if __name__ == "__main__":
    main()
