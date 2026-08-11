# -*- coding: utf-8 -*-
"""
Вырезать один месяц из годовых планов sro files/plany/*.docx
и сохранить отдельные Word-файлы в формате месячных планов
(как 001-plan-proverok_ogps_august_2026):
  справа «УТВЕРЖДАЮ» / Председатель КК / Д.В. Глебов / дата
  — НЕ как у годовых (УТВЕРЖДЕНО + протокол Правления).

Помеченные организации (исключена / проверка в 2026) —
цветной заливкой и причина в скобках в названии.

Пример:
  py export_month_plans_docx.py --month сентябрь
  py export_month_plans_docx.py --month сентябрь --skip-sro ОСОВС,ОСОТ,НОСО,ОСОЕС
"""
from __future__ import annotations

import argparse
import json
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from config_keys import SRO_FILES_DIR
from plan_proverok_export import (
    COL_WIDTHS_DXA,
    MONTH_FILE,
    SRO_META,
    set_cell_shading,
    set_cell_width,
    set_document_defaults,
    set_run_font,
    set_table_grid,
    setup_plan_row_numbering,
    write_cell_paragraphs,
    write_data_row,
)
from reestr_sync import plany_key_from_filename, sro_display_name

MONTHS = [
    "январь",
    "февраль",
    "март",
    "апрель",
    "май",
    "июнь",
    "июль",
    "август",
    "сентябрь",
    "октябрь",
    "ноябрь",
    "декабрь",
]

SKIP_DEFAULT = {"ОСОВС", "ОСОТ", "НОСО", "ОСОЕС"}
# исключена
RED_FILL = "FFC7CE"
RED_FONT = RGBColor(0x9C, 0x00, 0x06)
# 2+ проверки в 2026 — смотреть в базе (плановая или нет)
ORANGE_FILL = "FFEB9C"
ORANGE_FONT = RGBColor(0x9C, 0x57, 0x00)
# 1 проверка в 2026 — на заметку, не значит «убрать»
YELLOW_FILL = "FFF2CC"
YELLOW_FONT = RGBColor(0x7F, 0x60, 0x00)

MONTH_GENITIVE = {
    1: "января",
    2: "февраля",
    3: "марта",
    4: "апреля",
    5: "мая",
    6: "июня",
    7: "июля",
    8: "августа",
    9: "сентября",
    10: "октября",
    11: "ноября",
    12: "декабря",
}

MONTH_PREPOSITIONAL = {
    "ЯНВАРЬ": "январь",
    "ФЕВРАЛЬ": "февраль",
    "МАРТ": "март",
    "АПРЕЛЬ": "апрель",
    "МАЙ": "май",
    "ИЮНЬ": "июнь",
    "ИЮЛЬ": "июль",
    "АВГУСТ": "август",
    "СЕНТЯБРЬ": "сентябрь",
    "ОКТЯБРЬ": "октябрь",
    "НОЯБРЬ": "ноябрь",
    "ДЕКАБРЬ": "декабрь",
}


def _format_approve_date(d: date | None = None) -> str:
    """«26» июня 2026г. — как в эталоне august."""
    d = d or date.today()
    return f"«{d.day}» {MONTH_GENITIVE[d.month]} {d.year}г."


def build_month_document(
    meta: dict,
    year: int,
    month_upper: str,
    items: list[dict],
    approve_date: date | None = None,
) -> Document:
    """Месячный план: шапка «УТВЕРЖДАЮ» справа + таблица (как august-эталон)."""
    doc = Document()
    set_document_defaults(doc)

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.5)

    # Правый блок утверждения (НЕ годовой УТВЕРЖДЕНО + протокол)
    approve_lines = [
        ("«УТВЕРЖДАЮ»", True),
        ("Председатель Контрольного комитета", False),
        ("Д.В. Глебов", False),
        ("", False),
        ("________________________________", False),
        (_format_approve_date(approve_date), False),
    ]
    for text, bold in approve_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.left_indent = Cm(14)
        if text:
            run = p.add_run(text)
            set_run_font(run, size=12, bold=bold)

    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(0)
    gap.paragraph_format.space_after = Pt(6)

    month_name = MONTH_PREPOSITIONAL.get(month_upper, month_upper.lower())
    title_org = meta.get("title_org") or meta.get("folder") or ""
    title2 = f"{title_org} на {month_name} {year} года."

    # Заголовок — НАД таблицей (не строки таблицы)
    for text in (
        "План проверок членов саморегулируемой организации",
        title2,
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        run = p.add_run(text)
        set_run_font(run, size=12, bold=True)

    # Только шапка колонок + данные
    n_rows = 1 + len(items)
    table = doc.add_table(rows=max(n_rows, 2), cols=5)
    table.autofit = False
    table.allow_autofit = False
    set_table_grid(table, COL_WIDTHS_DXA)

    headers = [
        ("№\nп/п", True),
        ("ОПФ", False),
        ("Наименование организации", False),
        ("ИНН", False),
        ("Адрес", False),
    ]
    for i, (label, twoline) in enumerate(headers):
        cell = table.rows[0].cells[i]
        lines = label.split("\n") if twoline else [label]
        write_cell_paragraphs(
            cell,
            lines,
            bold=True,
            size=12,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        set_cell_width(cell, COL_WIDTHS_DXA[i])
        set_cell_shading(cell, None)

    row_numbering_id = setup_plan_row_numbering(doc)
    ridx = 1
    for item in items:
        write_data_row(
            table,
            ridx,
            [
                "",
                item.get("opf") or "",
                item.get("name") or "",
                item.get("inn") or "",
                item.get("address") or "",
            ],
            bold=False,
            size=12,
            center_cols=set(),
            row_numbering_id=row_numbering_id,
        )
        for i, w in enumerate(COL_WIDTHS_DXA):
            set_cell_width(table.rows[ridx].cells[i], w)
        ridx += 1

    tbl = table._tbl
    while len(table.rows) > ridx:
        tbl.remove(table.rows[-1]._tr)

    return doc


def _norm(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").replace("\xa0", " ")).strip()


def _month_label(month: str) -> str:
    m = month.strip().lower()
    for name in MONTHS:
        if name in m or m in name:
            return name.upper()
    raise ValueError(f"Неизвестный месяц: {month!r}")


def _unique_cells(row) -> list[str]:
    out: list[str] = []
    for cell in row.cells:
        txt = _norm(cell.text)
        if txt and txt not in out:
            out.append(txt)
    return out


def _read_header_meta(doc: Document, sro_label: str) -> dict:
    paras = [_norm(p.text) for p in doc.paragraphs if _norm(p.text)]
    decision = ""
    protocol = ""
    title_org = ""
    for t in paras:
        low = t.lower()
        if "решени" in low:
            decision = t
        elif "протокол" in low:
            protocol = t
    if doc.tables:
        title_row = _norm(doc.tables[0].rows[1].cells[0].text) if len(doc.tables[0].rows) > 1 else ""
        # «Ассоциации … на 2026 календарный год.»
        title_org = re.sub(r"\s+на\s+\d{4}.*$", "", title_row, flags=re.I).strip(" .")
    meta = SRO_META.get(sro_label, {}).copy()
    meta.setdefault("folder", sro_label)
    meta["decision"] = decision or meta.get("decision") or f"решением Правления СРО «{sro_label}»"
    meta["protocol"] = protocol or meta.get("protocol") or ""
    meta["title_org"] = title_org or meta.get("title_org") or sro_label
    return meta


def _parse_data_row(cells: list[str]) -> dict | None:
    inn = ""
    inn_index = -1
    for idx, text in enumerate(cells):
        clean = text.replace(" ", "")
        if clean.isdigit() and 9 <= len(clean) <= 12:
            inn = clean
            inn_index = idx
            break
    if not inn:
        return None

    before = [t for t in cells[:inn_index] if not t.replace(" ", "").isdigit()]
    # типично: ОПФ + название; иногда только название
    opf = ""
    name = "Организация СРО"
    if len(before) >= 2:
        # короткое ОПФ (ООО, АО, ИП …)
        cand = before[0]
        if len(cand) <= 12:
            opf = cand
            name = max(before[1:], key=len)
        else:
            name = max(before, key=len)
    elif before:
        name = before[0]

    addr = ""
    after = cells[inn_index + 1 :]
    if after:
        addr = max(after, key=len)

    return {"opf": opf, "name": name, "inn": inn, "address": addr}


def extract_month_from_plans(
    plany_dir: Path,
    month_target: str,
    skip_sro: set[str],
) -> dict[str, dict]:
    """sro_label -> {meta, items:[{opf,name,inn,address}]}"""
    month_upper = _month_label(month_target)
    result: dict[str, dict] = {}

    for file_path in sorted(plany_dir.glob("*.docx")):
        sro_key = plany_key_from_filename(file_path.name)
        sro_label = sro_display_name(sro_key)
        if sro_label in skip_sro or sro_key in skip_sro:
            continue

        doc = Document(str(file_path))
        meta = _read_header_meta(doc, sro_label)
        items: list[dict] = []
        current_month = ""

        for table in doc.tables:
            for row in table.rows:
                cells = _unique_cells(row)
                if not cells:
                    continue
                joined = " ".join(cells).lower()
                if len(cells) <= 3 and any(m in joined for m in MONTHS):
                    for m in MONTHS:
                        if m in joined:
                            current_month = m.upper()
                            break
                    continue
                if current_month != month_upper:
                    continue
                parsed = _parse_data_row(cells)
                if parsed:
                    items.append(parsed)

        if items:
            result[sro_label] = {"meta": meta, "items": items, "source": file_path.name}
    return result


def load_flag_map(path: Path | None) -> dict[tuple[str, str], dict]:
    """(sro, inn) -> {reason, mark}"""
    if not path or not path.exists():
        return {}
    data = json.loads(path.read_text(encoding="utf-8"))
    out: dict[tuple[str, str], dict] = {}
    for row in data:
        sro = row.get("sro") or ""
        inn = str(row.get("inn") or "").replace(" ", "")
        reason = row.get("reason") or ""
        if sro and inn and reason:
            out[(sro, inn)] = {
                "reason": reason,
                "mark": row.get("mark")
                or (
                    "duplicate"
                    if "задваива" in reason.lower()
                    else ("exclude" if "исключ" in reason else "one")
                ),
            }
    return out


def paint_flagged_rows(doc: Document, mark_by_inn: dict[str, str]) -> int:
    """Заливка строк: exclude=красн., multi=оранж., one=жёлт."""
    painted = 0
    if not doc.tables:
        return 0
    table = doc.tables[0]
    for row in table.rows:
        cells_text = [_norm(c.text) for c in row.cells]
        inn = ""
        for t in cells_text:
            clean = t.replace(" ", "")
            if clean.isdigit() and 9 <= len(clean) <= 12:
                inn = clean
                break
        mark = mark_by_inn.get(inn or "")
        if not mark:
            continue
        if mark in ("exclude", "duplicate"):
            fill, font = RED_FILL, RED_FONT
        elif mark == "multi":
            fill, font = ORANGE_FILL, ORANGE_FONT
        else:
            fill, font = YELLOW_FILL, YELLOW_FONT
        painted += 1
        for cell in row.cells:
            set_cell_shading(cell, fill)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = font
                    set_run_font(run, size=12, bold=(mark == "multi"))
    return painted


def apply_notes_to_items(
    items: list[dict], sro: str, flag_map: dict[tuple[str, str], dict]
) -> dict[str, str]:
    """Добавляет (причина) в название. Возвращает inn -> mark."""
    mark_by_inn: dict[str, str] = {}
    for item in items:
        info = flag_map.get((sro, item["inn"]))
        if not info:
            continue
        reason = info["reason"]
        note = f"({reason})"
        if note not in item["name"]:
            item["name"] = f"{item['name']} {note}"
        mark_by_inn[item["inn"]] = info["mark"]
    return mark_by_inn


def export_month(
    month: str,
    out_dir: Path,
    skip_sro: set[str],
    flag_path: Path | None,
    year: int = 2026,
    approve_date: date | None = None,
) -> list[str]:
    plany_dir = Path(SRO_FILES_DIR) / "plany"
    data = extract_month_from_plans(plany_dir, month, skip_sro)
    flag_map = load_flag_map(flag_path)
    month_upper = _month_label(month)
    month_file = MONTH_FILE[month_upper]
    out_dir.mkdir(parents=True, exist_ok=True)

    summary: list[str] = []
    for idx, (sro, bundle) in enumerate(sorted(data.items()), 1):
        items = bundle["items"]
        mark_by_inn = apply_notes_to_items(items, sro, flag_map)
        meta = bundle["meta"]
        doc = build_month_document(
            meta,
            year,
            month_upper,
            items,
            approve_date=approve_date,
        )
        painted = paint_flagged_rows(doc, mark_by_inn)

        safe = re.sub(r'[<>:"/\\|?*]+', "_", sro).strip()
        fname = f"{idx:03d}-plan-proverok_{safe}_{month_file}_{year}.docx"
        path = out_dir / fname
        doc.save(str(path))
        summary.append(
            f"{sro}: {len(items)} орг., помечено {len(mark_by_inn)} "
            f"(закрашено {painted}) -> {fname}"
        )
    (out_dir / "_отчет.txt").write_text("\n".join(summary), encoding="utf-8")
    return summary


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--month", default="сентябрь")
    ap.add_argument(
        "--out",
        default=str(Path(r"C:\Users\User\Desktop") / "Планы_сентябрь_2026_Word"),
    )
    ap.add_argument("--skip-sro", default=",".join(sorted(SKIP_DEFAULT)))
    ap.add_argument(
        "--flagged",
        default=str(Path(__file__).parent / "september_plan_check_for_word.json"),
        help="JSON с reason/mark (исключена / N проверок в 2026)",
    )
    ap.add_argument("--year", type=int, default=2026)
    ap.add_argument(
        "--approve-date",
        default=None,
        help="Дата утверждения ДД.ММ.ГГГГ (по умолчанию сегодня)",
    )
    args = ap.parse_args()

    skip = {p.strip() for p in args.skip_sro.split(",") if p.strip()}
    flag_path = Path(args.flagged) if args.flagged else None
    approve = None
    if args.approve_date:
        d, m, y = args.approve_date.split(".")
        approve = date(int(y), int(m), int(d))
    summary = export_month(
        args.month, Path(args.out), skip, flag_path, args.year, approve
    )
    print(f"OUT: {args.out}")
    for line in summary:
        print(line)
    print(f"Всего файлов: {len(summary)}")


if __name__ == "__main__":
    main()
