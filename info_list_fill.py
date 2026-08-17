"""Автозаполнение бланков из открытых данных реестра.

- Информационный лист — все 15 СРО («Местонахождение» → юр. адрес).
- Заявление о внесении изменений — выключено.
- Заявление на проверку — все 15 СРО.
- Доверенность — автозаполнение **выключено** (вёрстка/шрифт не устроили; код `build_doverennost_docx` оставлен, `DOVERENNOST_FILL_SRO_IDS` пустой).
"""

from __future__ import annotations

import os
import re
import tempfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from reestr_sync import (
    SRO_SOURCES,
    ensure_membership_details,
    get_org_memberships,
    sro_display_name,
)

# Информационный лист — все партнёрские СРО
INFO_LIST_FILL_SRO_IDS = frozenset(SRO_SOURCES.keys())
# Заявление об изменениях: автозаполнение выключено (пустое множество)
ZAYAVLENIE_FILL_SRO_IDS = frozenset()
# Заявление на проверку: все партнёрские СРО (шаблоны однотипные)
ZAYAVLENIE_PROVERKA_FILL_SRO_IDS = frozenset(SRO_SOURCES.keys())
# Доверенность: автозаполнение выключено (вернуть — заполнить DOVERENNOST_FILL_SRO_IDS)
DOVERENNOST_FILL_SRO_IDS = frozenset()

# Название Ассоциации в тексте доверенности (как в бланках сайтов)
DOVERENNOST_ASSOC_NAME: dict[str, str] = {
    "OGPS": "Объединение генеральных подрядчиков в строительстве",
    "GPS": "Генеральных подрядчиков в строительстве",
    "MOTS": "Межрегиональное объединение таврических строителей",
    "OGPP": "Объединение градостроительного планирования и проектирования",
    "OSO": "Объединение строительных организаций среднего и малого бизнеса",
    "SPROF": "Содружество профессиональных проектировщиков в строительстве",
    "PRIIS": "Профессионалы рынка инженерных изысканий в области строительства",
    "OPP": "Организаций профессионального проектирования",
    "NOSO": "Нижегородское объединение строительных организаций",
    "OSOES": "Объединение строительных организаций Екатеринбурга и Свердловской области",
    "OSOT": "Объединение строительных организаций Татарстана",
    "SOVS": "Объединение строительных организаций Восточной Сибири",
    "OGPO": "Объединение градостроительных проектных организаций",
    "MGEO": "Межрегиональное объединение изыскателей",
    "GEOIND": "Объединение изыскателей ГеоИндустрия",
}

# Совместимость со старыми импортами в bot_FINAL_GOLD
INFO_LIST_DISCLAIMER = ""
ZAYAVLENIE_IZM_DISCLAIMER = ""


def auto_fill_source_disclaimer(sro_id: str | None, *, doc_kind: str = "info_list") -> str:
    """Подпись к автозаполненному бланку: откуда данные + проверка перед подачей."""
    src = SRO_SOURCES.get(sro_id or "") or {}
    site = (src.get("list_url") or "").rstrip("/")
    if site.endswith("/reestr"):
        site_home = site[: -len("/reestr")] or site
    else:
        site_home = site or "сайта СРО"
    sro_name = src.get("name") or sro_display_name(sro_id or "") or "СРО"
    reestr_url = src.get("list_url") or site_home

    if doc_kind == "zayavlenie_izmeneniya":
        what = "наименование, адреса, ИНН, ОГРН и руководитель"
        extra = " Отметьте причину изменений."
    elif doc_kind == "zayavlenie_proverka":
        what = "наименование, ИНН, юр. адрес, рег. № и ФИО руководителя"
        extra = (
            " Ф.И.О. ответственного за документы, телефоны, email и дату проверки "
            "заполните вручную."
        )
    elif doc_kind == "doverennost":
        what = "наименование организации и должность/ФИО руководителя"
        extra = (
            " Дату выдачи, Ф.И.О. и паспортные данные доверенного лица, "
            "срок действия доверенности и подписи заполните вручную."
        )
    else:
        what = "рег. №, ИНН, юридический адрес (Местонахождение в реестре), руководитель и страхование"
        extra = (
            " Фактический адрес, телефоны, бухгалтер, ответственный и НРС подставляются из ответов в боте, если вы прошли опрос; иначе заполните вручную."
        )

    return (
        f"⚠️ <b>{sro_name}:</b> {what} взяты из официального реестра "
        f"<a href=\"{reestr_url}\">{reestr_url}</a> "
        f"(сайт СРО: <a href=\"{site_home}\">{site_home}</a>)."
        f"{extra} "
        "<b>Обязательно проверьте все сведения перед подачей</b> — "
        "данные могли измениться, часть полей заполните вручную."
    )


def collect_org_form_data(
    inn: str,
    plany_data: dict | None,
    reestr_data: dict | None,
    *,
    preferred_sro_id: str | None = None,
) -> dict | None:
    if not plany_data and not reestr_data:
        return None

    plany_data = plany_data or {}
    reestr_data = reestr_data or {}
    memberships = get_org_memberships(reestr_data)

    primary = None
    if preferred_sro_id and preferred_sro_id in memberships:
        primary = memberships[preferred_sro_id]
    if primary is None:
        for mem in memberships.values():
            if (mem.get("status") or "").startswith("Член"):
                primary = mem
                break
    if primary is None and memberships:
        primary = next(iter(memberships.values()))

    if primary:
        primary = ensure_membership_details(primary)
        if preferred_sro_id and preferred_sro_id in memberships:
            memberships[preferred_sro_id] = primary
        elif primary.get("sro_id") and primary["sro_id"] in memberships:
            memberships[primary["sro_id"]] = primary

    full_name = None
    short_name = reestr_data.get("title") or plany_data.get("name")
    reg_date = None
    reg_number = None
    location = None
    director = None
    insurance_company = None
    insurance_sum = None
    ogrn = None
    sro_name = None
    status = None
    sro_id = preferred_sro_id

    if primary:
        full_name = primary.get("full_name") or primary.get("title") or primary.get("short_name")
        short_name = primary.get("short_name") or primary.get("title") or short_name
        reg_date = primary.get("reg_date")
        reg_number = primary.get("reg_number")
        location = primary.get("location")
        director = primary.get("director")
        insurance_company = primary.get("insurance_company")
        insurance_sum = primary.get("insurance_sum")
        ogrn = primary.get("ogrn")
        sro_id = primary.get("sro_id") or preferred_sro_id
        sro_name = primary.get("sro_name") or sro_display_name(sro_id or "")
        status = primary.get("status")

    inn_digits = re.sub(r"\D", "", inn)
    full_name, short_name, director = _normalize_names_and_director(
        inn_digits,
        full_name=full_name,
        short_name=short_name,
        director=director,
    )

    return {
        "inn": inn_digits,
        "full_name": _clean_name(full_name or short_name),
        "short_name": _clean_name(short_name),
        "reg_date": reg_date or "",
        "reg_number": (reg_number or "").strip(),
        "location": _clean_address(location),
        "director": _clean_director(director),
        "insurance_company": _clean_insurance_company(insurance_company),
        "insurance_sum": _clean_insurance_sum(insurance_sum),
        "ogrn": re.sub(r"\D", "", ogrn or ""),
        "sro_name": sro_name or "",
        "sro_id": sro_id,
        "status": status or "",
        "is_ip": len(inn_digits) == 12,
    }


def _is_individual_entrepreneur(inn: str, full_name: str | None, short_name: str | None) -> bool:
    if len(inn) == 12:
        return True
    blob = f"{full_name or ''} {short_name or ''}".lower()
    return "индивидуальный предприниматель" in blob or blob.strip().startswith("ип ")


def _normalize_names_and_director(
    inn: str,
    *,
    full_name: str | None,
    short_name: str | None,
    director: str | None,
) -> tuple[str | None, str | None, str | None]:
    """
    Для ИП в реестре ФИО часто лежит в full_name, а «руководитель» пуст.
    Тогда: наименование = ИП + ФИО, руководитель = то же лицо.
    """
    full_name = (full_name or "").strip() or None
    short_name = (short_name or "").strip() or None
    director = (director or "").strip() or None

    if not _is_individual_entrepreneur(inn, full_name, short_name):
        return full_name, short_name, director

    fio_full = full_name or short_name or ""
    # Убрать уже стоящий префикс ИП, чтобы не задвоить
    for prefix in (
        "индивидуальный предприниматель",
        "ип ",
    ):
        low = fio_full.lower()
        if low.startswith(prefix):
            fio_full = fio_full[len(prefix) :].strip(" .")
            break

    fio_short = short_name or fio_full
    for prefix in ("ип ", "индивидуальный предприниматель"):
        if fio_short.lower().startswith(prefix):
            fio_short = fio_short[len(prefix) :].strip(" .")
            break

    if fio_full:
        full_name = f"Индивидуальный предприниматель {fio_full}"
    if fio_short:
        short_name = f"ИП {fio_short}"
    if not director and fio_full:
        director = f"Индивидуальный предприниматель {fio_full}"
    return full_name, short_name, director


def _clean_name(value: str | None) -> str:
    if not value:
        return ""
    return value.replace('"', "").replace("«", "").replace("»", "").strip()


def _clean_address(value: str | None) -> str:
    text = _clean_name(value)
    if not text:
        return ""
    low = text.lower()
    if "не предостав" in low or text == "—":
        return ""
    return text


def _clean_director(value: str | None) -> str:
    text = _clean_name(value)
    if not text:
        return ""
    low = text.lower()
    if "не предостав" in low:
        return ""
    text = re.sub(r"\s+", " ", text).strip()
    # «И.о.» / «ИО» → полностью «Исполняющий обязанности …»
    text = re.sub(
        r"^(?:и\.\s*о\.|ио)\s+генерального\s+директора\b",
        "Исполняющий обязанности Генерального директора",
        text,
        count=1,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r"^(?:и\.\s*о\.|ио)\s+директора\b",
        "Исполняющий обязанности директора",
        text,
        count=1,
        flags=re.IGNORECASE,
    )
    # Уже длинная форма — выровнять регистр должности
    text = re.sub(
        r"^исполняющ(?:ий|ая)\s+обязанности\s+генерального\s+директора\b",
        "Исполняющий обязанности Генерального директора",
        text,
        count=1,
        flags=re.IGNORECASE,
    )
    return text


def _clean_insurance_company(value: str | None) -> str:
    if not value:
        return ""
    # На сайте часто склеено: «…Гранта»Лицензия: …Местонахождение: …
    text = re.split(r"Лицензия\s*:", value, maxsplit=1, flags=re.IGNORECASE)[0]
    text = re.split(r"Местонахождение\s*:", text, maxsplit=1, flags=re.IGNORECASE)[0]
    text = _clean_name(text)
    if "не предостав" in text.lower():
        return ""
    return text


def _clean_insurance_sum(value: str | None) -> str:
    if not value:
        return ""
    text = value.replace("\xa0", " ").strip()
    if "не предостав" in text.lower():
        return ""
    # Оставить «5 000 000» / «5 000 000 руб.»
    return re.sub(r"\s+", " ", text).strip()


def _apply_run_font(
    run,
    *,
    bold: bool = False,
    font_name: str = "Times New Roman",
    size_pt: float = 12,
) -> None:
    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)


def _set_label_value_paragraph(
    paragraph,
    label: str,
    value: str,
    *,
    label_bold: bool = True,
    value_bold: bool = False,
    font_name: str = "Times New Roman",
) -> None:
    """Подпись жирная, значение после двоеточия — обычным шрифтом."""
    paragraph.clear()
    label_run = paragraph.add_run(f"{label.rstrip()}: ")
    _apply_run_font(label_run, bold=label_bold, font_name=font_name)
    value_run = paragraph.add_run(value)
    _apply_run_font(value_run, bold=value_bold, font_name=font_name)


def _fill_paragraph_after_label(
    paragraphs,
    label_part: str,
    value: str,
    *,
    bold_label: bool = False,
    font_name: str | None = None,
) -> bool:
    if not value:
        return False
    needle = label_part.lower()
    font = font_name or "Times New Roman"
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text
        if needle not in text.lower():
            continue
        stripped = text.rstrip()
        if ":" in stripped:
            head, _sep, _tail = stripped.partition(":")
            if bold_label or font_name:
                _set_label_value_paragraph(
                    paragraph,
                    head,
                    value,
                    label_bold=True,
                    value_bold=False,
                    font_name=font,
                )
            else:
                paragraph.text = f"{head}: {value}"
            return True
        if index + 1 < len(paragraphs):
            next_para = paragraphs[index + 1]
            if not next_para.text.strip() or len(next_para.text.strip()) < 3:
                next_para.text = value
                if font_name:
                    for run in next_para.runs:
                        _apply_run_font(run, bold=False, font_name=font)
                return True
    return False


def _clear_hint_paragraph_after(
    paragraphs, label_part: str, *, hint_substr: str
) -> None:
    """Убирает следующую строку-подсказку в скобках, чтобы не дублировать после заполнения."""
    needle = label_part.lower()
    hint = hint_substr.lower()
    for index, paragraph in enumerate(paragraphs):
        if needle not in paragraph.text.lower():
            continue
        if index + 1 >= len(paragraphs):
            return
        nxt = paragraphs[index + 1]
        t = nxt.text.strip()
        if t.startswith("(") and hint in t.lower() and len(t) < 80:
            nxt.text = ""
        return


def _fill_underscore_field(paragraphs, label_part: str, value: str) -> bool:
    """Подставляет значение вместо первого ряда подчёркиваний в строке с меткой."""
    if not value:
        return False
    needle = label_part.lower()
    for paragraph in paragraphs:
        text = paragraph.text
        if needle not in text.lower():
            continue
        if value in text:
            return True
        if re.search(r"_{3,}", text):
            paragraph.text = _replace_underscores_keep_width(text, value, count=1)
            return True
        # Нет подчёркиваний — дописать в конец строки
        paragraph.text = text.rstrip() + f" {value}"
        return True
    return False



def _il_write_para(
    paragraph,
    text: str,
    *,
    bold: bool = False,
    font_name: str = "Times New Roman",
    size_pt: float = 12,
) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    _apply_run_font(run, bold=bold, font_name=font_name, size_pt=size_pt)


def _is_placeholder_line(text: str) -> bool:
    t = (text or "").strip()
    if not t:
        return True
    if t.startswith("(") and len(t) < 90:
        return True
    if set(t.replace("\t", "")) <= set("._ "):
        return True
    return False


def _fill_label_then_blank_line(
    paragraphs,
    label_part: str,
    value: str,
    *,
    font_name: str = "Times New Roman",
) -> bool:
    """Подпись остаётся на своей строке, значение — на следующей пустой/подсказке.

    Так длинный адрес не наезжает на поле ниже.
    """
    if not value:
        return False
    needle = label_part.lower()
    for index, paragraph in enumerate(paragraphs):
        if needle not in (paragraph.text or "").lower():
            continue
        head = paragraph.text.split(":", 1)[0]
        head = re.sub(r"\s*\([^)]*\)", "", head)
        head = re.sub(r"\s+", " ", head).strip()
        _il_write_para(paragraph, f"{head}:", bold=True, font_name=font_name)
        for j in range(index + 1, min(index + 4, len(paragraphs))):
            nxt = paragraphs[j]
            raw = nxt.text or ""
            low = raw.lower()
            if any(
                k in low
                for k in (
                    "моб",
                    "тел.:",
                    "e-mail",
                    "идентификационн",
                    "главный бухгалтер",
                    "лицо, ответствен",
                    "наличие специалист",
                    "сведения о страхован",
                    "юридический адрес",
                    "адрес почтовый",
                    "адрес фактическ",
                    "сокращенное наименован",
                    "полное наименован",
                    "руководитель",
                )
            ) and not _is_placeholder_line(raw):
                break
            if _is_placeholder_line(raw):
                _il_write_para(nxt, value, bold=False, font_name=font_name)
                return True
        _set_label_value_paragraph(
            paragraph, head, value, label_bold=True, font_name=font_name
        )
        return True
    return False


def _replace_underscores_keep_width(text: str, value: str, count: int = 1) -> str:
    value = (value or "").strip()
    if not value:
        return text

    def _repl(match: re.Match[str]) -> str:
        width = len(match.group(0))
        body = f" {value} "
        if len(body) < width:
            body = body + (" " * (width - len(body)))
        return body

    return re.sub(r"_{3,}", _repl, text, count=count)


def _fill_phone_email_combo(paragraph, phone: str = "", email: str = "") -> bool:
    """Телефон слева, e-mail справа — табы из шаблона не съедаем."""
    raw = paragraph.text or ""
    phone = (phone or "").strip()
    email = (email or "").strip()
    if not phone and not email:
        return False
    tab_n = raw.count("\t")
    tabs = "\t" * (tab_n if tab_n >= 2 else 2)
    low = raw.lower()
    if "моб" in low:
        left = "моб. тел.: " + phone
        keep_bold = False
    else:
        left = "Тел.: " + phone
        keep_bold = True
    right = "e-mail – " + email
    paragraph.clear()
    run = paragraph.add_run(f"{left}{tabs}{right}")
    _apply_run_font(run, bold=keep_bold, font_name="Times New Roman")
    return True


def _fill_insurance_contract_line(
    paragraphs,
    number: str = "",
    date_from: str = "",
    date_to: str = "",
    amount: str = "",
) -> bool:
    number = (number or "").strip()
    date_from = (date_from or "").strip()
    date_to = (date_to or "").strip()
    amount = (amount or "").strip()
    if not (number or date_from or date_to or amount):
        return False
    for paragraph in paragraphs:
        text = paragraph.text or ""
        if "договор" not in text.lower() or "страхован" not in text.lower():
            continue
        orig = text
        if number:
            text = _replace_underscores_keep_width(text, number, count=1)
        if date_from:
            text = re.sub(
                r"(\sс\s)_{3,}",
                lambda m, v=date_from: m.group(1) + _replace_underscores_keep_width(m.group(0)[len(m.group(1)):], v),
                text,
                count=1,
            )
            if date_from not in text:
                text = re.sub(r"\sс\s_+", f" с {date_from} ", text, count=1)
        if date_to:
            if re.search(r"\sпо\s_+", text):
                text = re.sub(r"\sпо\s_+", f" по {date_to} ", text, count=1)
        if amount:
            text = re.sub(
                r"(на сумму\s*)_+(\s*руб\.?)?",
                lambda m, v=amount: f"{m.group(1)}{v} руб.",
                text,
                count=1,
                flags=re.IGNORECASE,
            )
        if text != orig:
            _il_write_para(paragraph, text, bold=False)
            return True
    return False


def _surname_initials(fio: str) -> str:
    parts = [p for p in (fio or "").split() if p]
    if not parts:
        return ""
    if len(parts) == 1:
        return parts[0]
    return parts[0] + " " + "".join(p[0] + "." for p in parts[1:] if p)


def _fill_signature_line(paragraphs, director: str) -> bool:
    if not director:
        return False
    position, fio = _split_director(director)
    initials = _surname_initials(fio or director)
    left = position or ""
    right = initials
    for paragraph in paragraphs:
        text = (paragraph.text or "").strip()
        if "страхован" in text.lower() or "сумм" in text.lower():
            continue
        if text.startswith("_") and len(text) > 20 and set(text) <= set("._ "):
            line = f"{left}\t\t\t{right}" if left else right
            _il_write_para(paragraph, line, bold=False)
            return True
    return False


def _fill_specialists_lines(paragraphs, specialists: str) -> bool:
    raw = (specialists or "").strip()
    if not raw:
        return False
    lines = [ln.strip(" ;") for ln in re.split(r"[\n;]+", raw) if ln.strip(" ;")]
    if not lines:
        return False
    for index, paragraph in enumerate(paragraphs):
        if "внесенных в нрс" not in (paragraph.text or "").lower():
            continue
        _il_write_para(
            paragraph,
            "Наличие специалистов, внесенных в НРС:",
            bold=True,
        )
        dest = index + 1
        placed = 0
        for line in lines:
            while dest < len(paragraphs):
                nxt = paragraphs[dest]
                low = (nxt.text or "").lower()
                if "сведения о страхован" in low:
                    dest = len(paragraphs)
                    break
                if _is_placeholder_line(nxt.text):
                    _il_write_para(nxt, line, bold=False)
                    dest += 1
                    placed += 1
                    break
                dest += 1
        return placed > 0
    return False


def _fill_info_list_user_fields(paragraphs, form_data: dict, filled: dict) -> None:
    """Поля опроса: на своих строках, без сдвига соседних полей."""
    fact = (form_data.get("fact_address") or "").strip()
    post = (form_data.get("post_address") or "").strip()
    if fact:
        filled["fact_address"] = _fill_label_then_blank_line(
            paragraphs, "Адрес фактического местонахождения", fact
        )
    if post:
        filled["post_address"] = _fill_label_then_blank_line(
            paragraphs, "Адрес почтовый", post
        )

    org_phone = (form_data.get("org_phone") or "").strip()
    org_email = (form_data.get("org_email") or "").strip()
    filled["org_contacts"] = False
    for paragraph in paragraphs:
        low = (paragraph.text or "").lower()
        if "тел" in low and "e-mail" in low and "моб" not in low:
            filled["org_contacts"] = _fill_phone_email_combo(
                paragraph, org_phone, org_email
            )
            break

    mob_lines = [
        p
        for p in paragraphs
        if "моб" in (p.text or "").lower() and "тел" in (p.text or "").lower()
    ]
    pairs = (
        (
            form_data.get("director_mobile") or "",
            form_data.get("director_email") or "",
        ),
        (
            form_data.get("accountant_phone") or "",
            form_data.get("accountant_email") or "",
        ),
        (
            form_data.get("responsible_phone") or "",
            form_data.get("responsible_email") or "",
        ),
    )
    filled["director_contacts"] = False
    filled["accountant_contacts"] = False
    filled["responsible_contacts"] = False
    flags = ("director_contacts", "accountant_contacts", "responsible_contacts")
    for index, paragraph in enumerate(mob_lines[:3]):
        phone, email = pairs[index]
        if phone or email:
            filled[flags[index]] = _fill_phone_email_combo(paragraph, phone, email)

    acc_fio = (form_data.get("accountant_fio") or "").strip()
    filled["accountant"] = False
    if acc_fio:
        filled["accountant"] = _fill_label_then_blank_line(
            paragraphs, "Главный бухгалтер", acc_fio
        )

    resp_bits = [
        (form_data.get("responsible_position") or "").strip(),
        (form_data.get("responsible_fio") or "").strip(),
    ]
    resp_text = ", ".join(b for b in resp_bits if b)
    filled["responsible"] = False
    if resp_text:
        filled["responsible"] = _fill_label_then_blank_line(
            paragraphs, "Лицо, ответственное за взаимодействие с СРО", resp_text
        )

    specialists = (form_data.get("specialists") or "").strip()
    filled["specialists"] = _fill_specialists_lines(paragraphs, specialists)

    filled["insurance_contract"] = _fill_insurance_contract_line(
        paragraphs,
        form_data.get("insurance_contract") or "",
        form_data.get("insurance_from") or "",
        form_data.get("insurance_to") or "",
        form_data.get("insurance_sum") or "",
    )
    filled["signature"] = _fill_signature_line(paragraphs, form_data.get("director") or "")


def _fill_insurance_sum_line(paragraphs, amount: str) -> bool:
    if not amount:
        return False
    for paragraph in paragraphs:
        text = paragraph.text
        if "на сумму" not in text.lower():
            continue
        if amount in text:
            return True
        new_text = re.sub(
            r"на сумму\s*_+\s*руб\.?",
            f"на сумму {amount} руб.",
            text,
            count=1,
            flags=re.IGNORECASE,
        )
        if new_text == text:
            new_text = re.sub(
                r"на сумму\s*_+",
                f"на сумму {amount} ",
                text,
                count=1,
                flags=re.IGNORECASE,
            )
        if new_text != text:
            paragraph.text = new_text
            return True
    return False


def _fill_header_reg_number(doc: Document, reg_number: str) -> bool:
    if not reg_number or not doc.paragraphs:
        return False
    header = doc.paragraphs[0].text
    if "Рег." not in header and "рег." not in header.lower():
        return False
    new_header = re.sub(
        r"(Рег\.\s*№)\s*_+",
        rf"\1 {reg_number}",
        header,
        count=1,
        flags=re.IGNORECASE,
    )
    if new_header == header:
        new_header = re.sub(
            r"(Рег\.\s*№)\s*",
            rf"\1 {reg_number} ",
            header,
            count=1,
            flags=re.IGNORECASE,
        )
    if new_header != header:
        doc.paragraphs[0].text = new_header
        return True
    return False


def _fill_inn_table(doc: Document, inn: str) -> bool:
    digits = re.sub(r"\D", "", inn)
    if not digits or not doc.tables:
        return False
    for table in doc.tables:
        if not table.rows:
            continue
        cells = table.rows[0].cells
        if len(cells) < 10:
            continue
        existing = "".join(c.text.strip() for c in cells)
        if existing and not all(ch.isdigit() or ch.isspace() for ch in existing):
            continue
        for i, cell in enumerate(cells):
            digit = digits[i] if i < len(digits) else ""
            cell.text = digit
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    _apply_run_font(run, bold=False, font_name="Times New Roman", size_pt=12)
        return True
    return False


def fill_info_list_docx(template_path: str, output_path: str, form_data: dict) -> dict:
    """Заполняет шаблон. Возвращает, что удалось подставить."""
    doc = Document(template_path)
    paragraphs = doc.paragraphs
    address = form_data.get("location", "")

    filled = {
        "reg_number": _fill_header_reg_number(doc, form_data.get("reg_number", "")),
        "inn_table": _fill_inn_table(doc, form_data.get("inn", "")),
        "full_name": _fill_label_then_blank_line(
            paragraphs,
            "Полное наименование юридического лица",
            form_data.get("full_name", ""),
        ),
        "short_name": _fill_label_then_blank_line(
            paragraphs,
            "Сокращенное наименование юридического",
            form_data.get("short_name", ""),
        ),
        "legal_address": _fill_label_then_blank_line(
            paragraphs,
            "Юридический адрес",
            address,
        ),
        "fact_address": False,
        "post_address": False,
        "director": _fill_label_then_blank_line(
            paragraphs,
            "Руководитель юридического лица",
            form_data.get("director", ""),
        ),
        "insurance_company": _fill_underscore_field(
            paragraphs,
            "Наименование Страховой компании",
            form_data.get("insurance_company", ""),
        ),
        "insurance_sum": False,
    }
    if not filled["inn_table"]:
        filled["inn_text"] = _fill_paragraph_after_label(
            paragraphs,
            "Идентификационный номер налогоплательщика",
            form_data.get("inn", ""),
        )
    else:
        filled["inn_text"] = False

    _fill_info_list_user_fields(paragraphs, form_data, filled)
    doc.save(output_path)
    return filled


def resolve_info_list_template(blanki_dir: str) -> str | None:
    """Для заполнения нужен .docx (python-docx)."""
    for name in ("info_list.docx", "info_list.doc"):
        path = os.path.join(blanki_dir, name)
        if os.path.isfile(path) and path.lower().endswith(".docx"):
            return path
    return None


def generate_info_list_for_inn(
    inn: str,
    blanki_dir: str,
    plany_data: dict | None,
    reestr_data: dict | None,
    *,
    preferred_sro_id: str | None = None,
    extra: dict | None = None,
) -> tuple[str | None, dict | None, dict | None]:
    """
    Returns (output_path, form_data, filled_flags) or (None, None, None).
    """
    sid = preferred_sro_id or ""
    if sid not in INFO_LIST_FILL_SRO_IDS:
        return None, None, None

    template_path = resolve_info_list_template(blanki_dir)
    if not template_path:
        return None, None, None

    form_data = collect_org_form_data(
        inn, plany_data, reestr_data, preferred_sro_id=sid
    )
    if not form_data:
        return None, None, None

    if extra:
        from info_list_quiz import resolve_quiz_addresses

        merged = resolve_quiz_addresses(
            form_data, {k: v for k, v in extra.items() if v}
        )
        for key, value in merged.items():
            if value:
                form_data[key] = value

    temp_dir = tempfile.gettempdir()
    safe_inn = form_data["inn"] or "org"
    output_path = os.path.join(temp_dir, f"info_list_{sid}_{safe_inn}.docx")
    filled = fill_info_list_docx(template_path, output_path, form_data)
    return output_path, form_data, filled


def _split_director(director: str) -> tuple[str, str]:
    text = (director or "").strip()
    if not text:
        return "", ""
    titles = (
        "Генеральный директор",
        "Исполнительный директор",
        "Коммерческий директор",
        "Индивидуальный предприниматель",
        "Управляющий — индивидуальный предприниматель",
        "Управляющий",
        "Президент",
        "Директор",
    )
    low = text.lower()
    for title in titles:
        if low.startswith(title.lower()):
            return title, text[len(title) :].strip(" ,–—-")
    return "", text


def _fill_single_cell_table(table, value: str) -> bool:
    if not value or not table.rows:
        return False
    cell = table.rows[0].cells[0]
    cell.text = value
    return True


def _fill_labeled_digit_row(table, label: str, digits: str) -> bool:
    if not digits or not table.rows:
        return False
    cells = table.rows[0].cells
    head = (cells[0].text or "").strip().upper().replace(" ", "")
    if not head.startswith(label.upper().replace(" ", "")):
        return False
    for i, cell in enumerate(cells[1:]):
        cell.text = digits[i] if i < len(digits) else ""
    return True


def _fill_director_table(table, director: str) -> bool:
    if not director or not table.rows:
        return False
    cells = table.rows[0].cells
    if len(cells) < 2:
        cells[0].text = director
        return True
    position, fio = _split_director(director)
    cells[0].text = position or director
    cells[1].text = fio if position else ""
    return True


def fill_zayavlenie_izmeneniya_docx(
    template_path: str, output_path: str, form_data: dict
) -> dict:
    doc = Document(template_path)
    tables = doc.tables
    address = form_data.get("location", "")
    filled = {
        "full_name": False,
        "short_name": False,
        "legal_address": False,
        "post_address": False,
        "fact_address": False,
        "inn": False,
        "ogrn": False,
        "director": False,
    }
    if len(tables) > 0:
        filled["full_name"] = _fill_single_cell_table(tables[0], form_data.get("full_name", ""))
    if len(tables) > 1:
        filled["short_name"] = _fill_single_cell_table(tables[1], form_data.get("short_name", ""))
    if len(tables) > 2:
        filled["legal_address"] = _fill_single_cell_table(tables[2], address)
    if len(tables) > 3:
        filled["post_address"] = _fill_single_cell_table(tables[3], address)
    if len(tables) > 4:
        filled["fact_address"] = _fill_single_cell_table(tables[4], address)
    if len(tables) > 6:
        filled["inn"] = _fill_labeled_digit_row(tables[6], "ИНН", form_data.get("inn", ""))
    if len(tables) > 7:
        filled["ogrn"] = _fill_labeled_digit_row(tables[7], "ОГРН", form_data.get("ogrn", ""))
    if len(tables) > 11:
        filled["director"] = _fill_director_table(tables[11], form_data.get("director", ""))

    doc.save(output_path)
    return filled


def generate_zayavlenie_izmeneniya_for_inn(
    inn: str,
    blanki_dir: str,
    plany_data: dict | None,
    reestr_data: dict | None,
    *,
    preferred_sro_id: str | None = "SPROF",
) -> tuple[str | None, dict | None, dict | None]:
    sid = preferred_sro_id or ""
    if sid not in ZAYAVLENIE_FILL_SRO_IDS:
        return None, None, None

    template_path = os.path.join(blanki_dir, "zayavlenie_izmeneniya.docx")
    if not os.path.isfile(template_path):
        return None, None, None

    form_data = collect_org_form_data(
        inn, plany_data, reestr_data, preferred_sro_id=sid
    )
    if not form_data:
        return None, None, None

    temp_dir = tempfile.gettempdir()
    safe_inn = form_data["inn"] or "org"
    output_path = os.path.join(temp_dir, f"zayavlenie_izmeneniya_{sid}_{safe_inn}.docx")
    filled = fill_zayavlenie_izmeneniya_docx(template_path, output_path, form_data)
    return output_path, form_data, filled


def _opf_and_quoted_name(form_data: dict) -> tuple[str, str]:
    """ОПФ для шаблона («ООО»/«ИП»/…) и наименование внутрь кавычек."""
    short = (form_data.get("short_name") or form_data.get("full_name") or "").strip()
    full = (form_data.get("full_name") or short).strip()
    if form_data.get("is_ip"):
        name = short or full
        name = re.sub(
            r"^(?:ИП|Индивидуальный предприниматель)\s+",
            "",
            name,
            flags=re.IGNORECASE,
        ).strip(" «»\"'")
        return "ИП", name

    blob = short or full
    opf = "ООО"
    for candidate in ("ПАО", "НАО", "АО", "ЗАО", "ООО", "НКО"):
        if re.match(rf"^{candidate}\b", blob, flags=re.IGNORECASE):
            opf = candidate
            blob = re.sub(rf"^{candidate}\s+", "", blob, count=1, flags=re.IGNORECASE)
            break
    name = blob.strip(" «»\"'")
    return opf, name


def _set_paragraph_plain(paragraph, text: str, *, size_pt: float = 12) -> None:
    """Перезаписать абзац одним run (Times New Roman) — шаблон простой, без сложной вёрстки."""
    # python-docx: очистить существующие runs
    for run in list(paragraph.runs):
        run.text = ""
    if paragraph.runs:
        run = paragraph.runs[0]
        run.text = text
    else:
        run = paragraph.add_run(text)
    run.bold = False
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")
    # убрать лишние пустые runs после первого
    for run in list(paragraph.runs[1:]):
        run._element.getparent().remove(run._element)

def fill_zayavlenie_proverka_docx(template_path: str, output_path: str, form_data: dict) -> dict:
    """
    Заполнить открытые поля заявления на проверку (шаблон ОГПС).
    Не трогаем: дату проверки, ФИО ответственного за документы, телефоны, email, подпись.
    """
    doc = Document(template_path)
    opf, quoted = _opf_and_quoted_name(form_data)
    inn = form_data.get("inn") or ""
    address = form_data.get("location") or ""
    reg_number = form_data.get("reg_number") or ""
    position, fio = _split_director(form_data.get("director") or "")
    if not position:
        position = "Индивидуальный предприниматель" if form_data.get("is_ip") else "Генеральный директор"
    if not fio and form_data.get("director"):
        fio = form_data["director"]

    filled = {
        "org_block": False,
        "responsible_org": False,
        "director_sign": False,
    }

    # Абзац с ООО / ИНН / адресом / рег. №
    for p in doc.paragraphs:
        text = p.text or ""
        if "просит рассмотреть вопрос о проведении" not in text:
            continue
        if "ООО «" not in text and "ИП «" not in text:
            # шаблон начинается с ООО
            if "«________________»" not in text and "ИНН" not in text:
                continue
        new = text
        # ОПФ в начале (шаблон: «ООО «…»)
        new = re.sub(r"^\s*ООО\s+", f"      {opf} ", new, count=1)
        new = re.sub(r"ООО\s*«_{5,}»", f"{opf} «{quoted}»", new, count=1)
        new = re.sub(r"«_{5,}»", f"«{quoted}»", new, count=1)
        if inn:
            new = re.sub(r"\(ИНН_{3,}\)", f"(ИНН {inn})", new, count=1)
            new = re.sub(r"ИНН_{3,}", f"ИНН {inn}", new, count=1)
        if address:
            new = re.sub(
                r"юридический адрес:_{3,}",
                f"юридический адрес:{address}",
                new,
                count=1,
            )
        if reg_number:
            new = re.sub(
                r"регистрационный № _{2,}",
                f"регистрационный № {reg_number}",
                new,
                count=1,
            )
            new = re.sub(
                r"регистрационный №_{2,}",
                f"регистрационный № {reg_number}",
                new,
                count=1,
            )
        if new != text:
            _set_paragraph_plain(p, new)
            filled["org_block"] = True
        break

    # Ответственный / представитель — только название организации (ФИО вручную)
    for p in doc.paragraphs:
        text = p.text or ""
        if "Ответственным от" not in text:
            continue
        new = text
        # Две разные длины подчёркиваний в шаблоне
        new = re.sub(r"от ООО «_{5,}»", f"от {opf} «{quoted}»", new, count=1)
        new = re.sub(
            r"представитель ООО «_{5,}»",
            f"представитель {opf} «{quoted}»",
            new,
            count=1,
        )
        if new != text:
            _set_paragraph_plain(p, new)
            filled["responsible_org"] = True
        break

    # Строка подписи руководителя
    for p in doc.paragraphs:
        text = p.text or ""
        if "(подпись)" not in text or "ФИО" not in text.upper().replace(".", ""):
            # шаблон: (ФИО)
            if "(подпись)" not in text:
                continue
            if "ФИО" not in text and "(ФИО)" not in text:
                continue
        if not fio:
            break
        # Заменить должность в начале, если шаблон «Генеральный директор»
        new = text
        new = re.sub(
            r"^\s*Генеральный директор\b",
            position,
            new,
            count=1,
        )
        new = re.sub(r"\(ФИО\)\s*", f"{fio}    ", new, count=1)
        if new != text:
            _set_paragraph_plain(p, new)
            filled["director_sign"] = True
        break

    doc.save(output_path)
    return filled


def generate_zayavlenie_proverka_for_inn(
    inn: str,
    blanki_dir: str,
    plany_data: dict | None,
    reestr_data: dict | None,
    *,
    preferred_sro_id: str | None = "OGPS",
) -> tuple[str | None, dict | None, dict | None]:
    sid = preferred_sro_id or ""
    if sid not in ZAYAVLENIE_PROVERKA_FILL_SRO_IDS:
        return None, None, None

    template_path = os.path.join(blanki_dir, "zayavlenie_proverka.docx")
    if not os.path.isfile(template_path):
        return None, None, None

    form_data = collect_org_form_data(
        inn, plany_data, reestr_data, preferred_sro_id=sid
    )
    if not form_data:
        return None, None, None

    temp_dir = tempfile.gettempdir()
    safe_inn = form_data["inn"] or "org"
    output_path = os.path.join(temp_dir, f"zayavlenie_proverka_{sid}_{safe_inn}.docx")
    filled = fill_zayavlenie_proverka_docx(template_path, output_path, form_data)
    return output_path, form_data, filled


def _fio_to_initials(fio: str) -> str:
    parts = [p for p in (fio or "").split() if p]
    if len(parts) >= 3:
        return f"{parts[0]} {parts[1][0]}.{parts[2][0]}."
    if len(parts) == 2:
        return f"{parts[0]} {parts[1][0]}."
    return fio or ""


def _dov_font(run, size_pt: float = 14, *, bold: bool = False, underline: bool = False) -> None:
    run.bold = bold
    run.underline = underline
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _dov_para_border_bottom(paragraph) -> None:
    """Нижняя граница абзаца = линия поля на всю ширину страницы."""
    from docx.oxml import OxmlElement

    pPr = paragraph._p.get_or_add_pPr()
    # убрать старую рамку, если есть
    for child in list(pPr):
        if child.tag == qn("w:pBdr"):
            pPr.remove(child)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_dov_para(
    doc: Document,
    text: str,
    *,
    bold: bool = False,
    center: bool = False,
    justify: bool = False,
    size_pt: float = 14,
    space_after_pt: float = 0,
    space_before_pt: float = 0,
    line_spacing: float = 1.15,
):
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_after = Pt(space_after_pt)
    pf.space_before = Pt(space_before_pt)
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    run = p.add_run(text)
    _dov_font(run, size_pt, bold=bold)
    return p


def _add_dov_hint(doc: Document, text: str) -> None:
    """Подпись под полем — мелким шрифтом по центру."""
    _add_dov_para(
        doc,
        text,
        center=True,
        size_pt=10,
        space_after_pt=8,
        space_before_pt=0,
        line_spacing=1.0,
    )


def _add_dov_field_line(
    doc: Document,
    text: str,
    hint: str,
    *,
    suffix: str = "",
    size_pt: float = 14,
) -> None:
    """
    Поле на всю ширину: текст на линии (нижняя граница абзаца) + подпись под ним.
    suffix — текст справа без подчёркивания (например « в лице»).
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(6)
    pf.line_spacing = 1.15
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    value = (text or "").strip() or " "
    run = p.add_run(value)
    _dov_font(run, size_pt, underline=False)
    if suffix:
        run2 = p.add_run(suffix)
        _dov_font(run2, size_pt, underline=False)
    _dov_para_border_bottom(p)
    if hint:
        _add_dov_hint(doc, hint)


def build_doverennost_docx(output_path: str, form_data: dict, sro_id: str) -> dict:
    """
    Собрать доверенность (.docx) как официальный бланк:
    Times New Roman 14, поля на всю ширину строки, подписи под линиями.
    Из реестра: организация и руководитель. Вручную: дата, доверенное лицо, срок, подписи.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.shared import Cm

    opf, quoted = _opf_and_quoted_name(form_data)
    org_line = f"{opf} «{quoted}»" if quoted else ""
    position, fio = _split_director(form_data.get("director") or "")
    if not position:
        position = (
            "Индивидуальный предприниматель"
            if form_data.get("is_ip")
            else "Генеральный директор"
        )
    if not fio and form_data.get("director"):
        fio = form_data["director"]
    # В бланке: должность и ФИО на одной строке (как в официальном шаблоне СРО)
    director_line = f"{position} {fio}".strip() if (position or fio) else ""
    initials = _fio_to_initials(fio) if fio else ""
    assoc = DOVERENNOST_ASSOC_NAME.get(sro_id) or (
        form_data.get("sro_name") or sro_display_name(sro_id) or ""
    )

    filled = {
        "org_name": bool(quoted),
        "director": bool(fio),
        "sign_block": bool(fio or position),
    }

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    try:
        normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    except Exception:
        pass

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.5)

    _add_dov_para(
        doc,
        "ДОВЕРЕННОСТЬ",
        bold=True,
        center=True,
        size_pt=16,
        space_after_pt=18,
        line_spacing=1.0,
    )

    # Дата — по центру, линия на ширину короткого поля
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_after = Pt(0)
    p_date.paragraph_format.line_spacing = 1.15
    r = p_date.add_run("«_____» ____________________ 20_____ г.")
    _dov_font(r, 14)
    _dov_para_border_bottom(p_date)
    _add_dov_hint(doc, "(дата выдачи доверенности)")

    # Организация на всю ширину (как в заполненных бланках)
    _add_dov_field_line(
        doc,
        org_line,
        "(наименование организации заявителя)",
    )
    _add_dov_para(
        doc,
        "в лице",
        size_pt=14,
        space_before_pt=4,
        space_after_pt=0,
        line_spacing=1.0,
    )

    _add_dov_field_line(
        doc,
        director_line,
        "(должность, фамилия, имя, отчество руководителя заявителя)",
    )

    _add_dov_para(
        doc,
        "действующего на основании Устава, уполномочивает",
        justify=True,
        size_pt=14,
        space_before_pt=10,
        space_after_pt=10,
    )

    _add_dov_field_line(
        doc,
        "",
        "(фамилия, имя, отчество, паспортные данные доверенного лица)",
    )
    # Вторая пустая линия под длинные паспортные данные (как в живых доверенностях)
    _add_dov_field_line(doc, "", "")

    _add_dov_para(
        doc,
        "представлять интересы организации при проведении контрольных проверок "
        f"СРО Ассоциацией {assoc},",
        justify=True,
        size_pt=14,
        space_before_pt=8,
        space_after_pt=0,
    )
    _add_dov_hint(doc, "(название Ассоциации)")

    _add_dov_para(
        doc,
        "с правом подписи соответствующих документов, предоставляемых в письменном "
        "виде при проведении указанных контрольных мероприятий, заверения своей "
        "подписью копий всех предоставляемых документов, а также совершения всех "
        "иных действий, связанных с выполнением настоящего поручения.",
        justify=True,
        size_pt=14,
        space_after_pt=12,
    )

    # Срок: текст + пустое поле на линии
    p_term = doc.add_paragraph()
    p_term.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_term.paragraph_format.space_after = Pt(0)
    p_term.paragraph_format.line_spacing = 1.15
    r1 = p_term.add_run("Настоящая доверенность выдана сроком на ")
    _dov_font(r1, 14)
    r2 = p_term.add_run("______________________________")
    _dov_font(r2, 14)
    r3 = p_term.add_run(",")
    _dov_font(r3, 14)
    _add_dov_hint(doc, "(указывается срок действия доверенности)")

    _add_dov_para(
        doc,
        "без права передоверия полномочий третьим лицам.",
        justify=True,
        size_pt=14,
        space_after_pt=18,
    )

    p_proxy = doc.add_paragraph()
    p_proxy.paragraph_format.space_after = Pt(0)
    p_proxy.paragraph_format.line_spacing = 1.15
    for part, und in (
        ("Подпись доверенного лица ", False),
        ("____________________", True),
        ("   ", False),
        ("____________________", True),
        (" удостоверяю.", False),
    ):
        rr = p_proxy.add_run(part)
        _dov_font(rr, 14, underline=und)
    _add_dov_hint(
        doc,
        "(Указать Ф. И.О.)                              (подпись доверенного лица)",
    )

    # Блок подписи руководителя — 3 колонки на всю ширину
    _add_dov_para(doc, "", space_after_pt=12)
    table = doc.add_table(rows=2, cols=3)
    table.autofit = True
    try:
        table.allow_autofit = True
    except Exception:
        pass

    cells0 = table.rows[0].cells
    cells1 = table.rows[1].cells
    cells0[0].text = ""
    cells0[1].text = ""
    cells0[2].text = ""

    def _cell_line(cell, text: str, hint: str) -> None:
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text or " ")
        _dov_font(run, 14)
        _dov_para_border_bottom(p)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        r = p2.add_run(hint)
        _dov_font(r, 9)

    _cell_line(cells0[0], position, "(должность руководителя)")
    _cell_line(cells0[1], "", "(подпись)")
    _cell_line(cells0[2], initials or " ", "(фамилия и инициалы)")
    # вторая строка таблицы не нужна — очистить
    for c in cells1:
        c.text = ""

    # убрать границы таблицы, оставить только линии полей в ячейках
    from docx.oxml import OxmlElement

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)

    _add_dov_para(doc, "М.П.", center=True, size_pt=14, space_before_pt=18, space_after_pt=0)

    doc.save(output_path)
    return filled


def generate_doverennost_for_inn(
    inn: str,
    blanki_dir: str,
    plany_data: dict | None,
    reestr_data: dict | None,
    *,
    preferred_sro_id: str | None = None,
) -> tuple[str | None, dict | None, dict | None]:
    """
    Returns (output_path, form_data, filled_flags) or (None, None, None).
    blanki_dir не обязателен для сборки (текст собираем в коде); проверяем, что
    у СРО вообще есть комплект бланков.
    """
    sid = preferred_sro_id or ""
    if sid not in DOVERENNOST_FILL_SRO_IDS:
        return None, None, None
    if blanki_dir and not os.path.isdir(blanki_dir):
        return None, None, None

    form_data = collect_org_form_data(
        inn, plany_data, reestr_data, preferred_sro_id=sid
    )
    if not form_data:
        return None, None, None

    temp_dir = tempfile.gettempdir()
    safe_inn = form_data["inn"] or "org"
    output_path = os.path.join(temp_dir, f"doverennost_{sid}_{safe_inn}.docx")
    filled = build_doverennost_docx(output_path, form_data, sid)
    return output_path, form_data, filled
