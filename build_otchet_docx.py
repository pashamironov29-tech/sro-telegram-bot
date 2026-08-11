# -*- coding: utf-8 -*-
"""Сборка единого отчёта DOCX для руководства. Запуск: python build_otchet_docx.py"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn

from ai_assistant import KEYWORD_RULES, SITE_TOPICS
from voprosy_faq import VOPROSY_ITEMS
from sro_site_qa import SRO_SITE_QA_ITEMS

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "OTCHET_KURATOR_I_BOT_SRO_GEN.docx"
BOT_VERSION = "2026-07-22-v50"


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_title_page(doc: Document) -> None:
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(
        "Telegram-бот СРО Ассоциации «ГЕН»\n"
        "Отчёт для руководства"
    )
    run.bold = True
    run.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Роль куратора · бот · база ответов · регионы и часовые пояса · "
        "перспектива EN-версии\n"
        f"Версия бота: {BOT_VERSION}    Дата: {date.today().strftime('%d.%m.%Y')}"
    )
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()
    p = doc.add_paragraph(
        "Документ объединяет описание цифрового сервиса для членов СРО, "
        "разграничение зон ответственности куратора и IT, а также количественную "
        "оценку встроенной базы ответов, ценность для регионов (в т.ч. Дальний Восток) "
        "и краткое предложение по англоязычной версии. Подготовлен для согласования "
        "с руководством и передачи в IT при развёртывании на сервере."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_page_break()


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()


def part_curator(doc: Document) -> None:
    heading(doc, "Часть I. Роль куратора бота", 1)

    heading(doc, "Для руководства — на одной странице", 2)
    bullet(doc, "Куратор — не программист и не IT-штат: владелец сценариев и качества для пользователя.")
    bullet(doc, "Доработки: формулировка задач + ИИ-инструменты + проверка в Telegram по скринам.")
    bullet(doc, "Зона куратора: что бот говорит, куда ведёт, актуальность FAQ, приёмка обновлений.")
    bullet(doc, "Не зона куратора: сервер 24/7, архитектура кода, юридические консультации, uptime без IT.")

    heading(doc, "Как назвать роль", 2)
    para(doc, "Рекомендуемое название: куратор / product-owner официального бота Ассоциации в Telegram.")

    heading(doc, "Что уже сделано (ценность для Ассоциации)", 2)
    add_table(
        doc,
        ["Направление", "Результат"],
        [
            ["Поиск организаций", "ИНН / название → карточка, план проверок, скачивание бланков"],
            ["FAQ и меню", "Структура как на сайте: вступление, членам, НОК/НРС, партнёры"],
            ["ИИ-помощник", "Подбор раздела srogen.ru + «Вопрос-ответ» без выдумывания норм"],
            ["Справочник", "Поиск по ФИО в «Контактах отделов» (пароль для членов)"],
            ["Новые темы", "Устав/документы, СРО «ГЕН» vs партнёры, законы, техрегулирование"],
            ["Качество", "Версии, чек-листы, разведение ИИ ↔ справочник ↔ реестр"],
        ],
    )

    heading(doc, "Куратор и программист — различие", 2)
    add_table(
        doc,
        ["Куратор", "Программист / IT"],
        [
            ["Процессы СРО, ожидания пользователей", "Код, архитектура, CI/CD"],
            ["Задачи словами, проверка в Telegram", "VPS, мониторинг, бэкапы"],
            ["Доработки с ИИ, ответственность за продукт", "Инфраструктура по ТЗ"],
        ],
    )

    heading(doc, "Риски для руководства", 2)
    bullet(doc, "Один процесс бота на один токен Telegram (иначе ошибка 409).")
    bullet(doc, "Справочник — только с паролем; утечка пароля = риск для внутренних данных.")
    bullet(doc, "ИИ не заменяет специалиста; спорные кейсы — телефон +7 (495) 775-81-11.")

    quote = doc.add_paragraph()
    quote.paragraph_format.left_indent = Cm(1)
    quote.add_run(
        "«Куратор бота — это не IT-отдел: это человек от Ассоциации, который с помощью ИИ "
        "поддерживает сервис в Telegram, чтобы член за минуту получил бланк, ссылку или контакт. "
        "Сложные вопросы — к специалистам СРО.»"
    ).italic = True


def part_bot(doc: Document) -> None:
    heading(doc, "Часть II. Презентация бота — что, куда и зачем", 1)

    heading(doc, "Зачем нужен бот", 2)
    para(
        doc,
        "Сайт srogen.ru содержит полную информацию, но членам СРО нужен быстрый ответ без "
        "обхода десятков разделов. Бот — цифровая витрина Ассоциации в Telegram: не заменяет "
        "юриста, а ведёт к официальным материалам и снимает рутину с секретариата.",
    )
    add_table(
        doc,
        ["Задача", "Как решает бот"],
        [
            ["Когда проверка у организации", "Поиск по ИНН → карточка + месяц из плана проверок"],
            ["Скачать бланк", "7 форм после карточки + архив на сайте"],
            ["Взносы, НОК, НРС", "FAQ, ИИ-помощник, ссылки на разделы"],
            ["Контакт отдела", "Контакты отделов + поиск по ФИО (с паролем)"],
            ["Партнёрское СРО / НО", "Партнёры и НО с поиском по названию"],
        ],
    )

    heading(doc, "Для кого", 2)
    for line in [
        "Действующие члены СРО — проверки, бланки, выписки, реестр.",
        "Вступающие — взносы, документы, сроки, специалисты.",
        "Кураторы и специалисты — НОК, НРС, кураторы НРС.",
        "Регионы и партнёры — филиалы, партнёрские СРО, НОСТРОЙ/НОПРИЗ.",
        "Сотрудники Ассоциации — меньше однотипных звонков.",
    ]:
        bullet(doc, line)

    heading(doc, "Главное меню — четыре опоры", 2)

    blocks = [
        (
            "Поиск организации",
            "ИНН или название → карточка из реестров 15+ партнёрских СРО, месяц плановой проверки "
            "из файлов plany/, кнопки скачивания 7 бланков. Команда /search.",
        ),
        (
            "Полезная информация",
            "Дерево FAQ (вступающим, членам, НОК/НРС, об Ассоциации, филиалы, партнёры). "
            "Проверяемые документы. Команда /info.",
        ),
        (
            "Контакты отделов",
            "Телефонный справочник отделов; пароль в config_keys — доступ для членов, "
            "поиск по фамилии и «телефон …».",
        ),
        (
            "ИИ-помощник",
            "Свободный текст → ключевые слова, FAQ, блоки сайта, при необходимости Groq. "
            "Приоритет: темы сайта и FAQ, затем реестр, если запрос похож на название организации.",
        ),
    ]
    for title, body in blocks:
        heading(doc, title, 3)
        para(doc, body)

    heading(doc, "Документы после поиска по ИНН", 2)
    forms = [
        "Информационный лист",
        "Заявление о внесении изменений",
        "Заявление на проверку",
        "Форма доверенности",
        "Сведения о специалистах",
        "Положения о контроле",
        "Уведомление ОДО",
    ]
    for i, f in enumerate(forms, 1):
        bullet(doc, f"{i}. {f}")

    heading(doc, "Сценарий демонстрации (5–7 минут)", 2)
    demo = [
        "/start — обзор возможностей",
        "ИНН организации — карточка, план, бланк",
        "Поиск по части названия",
        "Полезная информация → расписание проверок",
        "ИИ: «нок», «размеры взносов», «где устав»",
        "Партнёры: «нострой»",
        "Контакты (пароль при необходимости)",
    ]
    for step in demo:
        bullet(doc, step)

    heading(doc, "Развитие после пилота", 2)
    bullet(doc, "VPS 24/7 (инструкция IT_ZAYAVKA_DEPLOY.md, VPS_SETUP.md).")
    bullet(doc, "Автообновление reestr_cache.json по расписанию.")
    bullet(doc, "Опционально: уведомления о месяце проверки, мессенджер MAX.")


def part_baza(doc: Document) -> None:
    heading(doc, "Часть III. База ответов и ИИ-помощник", 1)

    n_kw = sum(len(p) for _, p in KEYWORD_RULES)
    add_table(
        doc,
        ["Показатель", "Количество", "Источник"],
        [
            ["Готовые Q&A (краткий ответ + ссылка)", "48", "voprosy_faq.py → srogen.ru/voprosy/"],
            ["Блоки навигации по сайту", "4", "sro_site_qa.py"],
            ["Итого карточек FAQ/сайт", "52", ""],
            ["Разделы сайта (URL + описание)", str(len(SITE_TOPICS)), "SITE_TOPICS"],
            ["Темы быстрого поиска", str(len(KEYWORD_RULES)), "KEYWORD_RULES"],
            ["Фраз-триггеров", f"~{n_kw}", "варианты вопросов пользователя"],
            ["Готовые тексты по кнопкам меню", "~15", "bot_FINAL_GOLD.py"],
        ],
    )

    para(
        doc,
        "52 и 39 — не сумма уникальных юридических текстов: одна тема (например, возврат взноса) "
        "может открываться из FAQ, меню и ключевых слов. Groq не хранит отдельную базу — "
        "выбирает раздел из каталога SITE_TOPICS или сообщает, что точного ответа нет.",
    )

    heading(doc, "Четыре блока навигации по сайту", 2)
    for item in SRO_SITE_QA_ITEMS:
        bullet(doc, item["label"])

    heading(doc, "Разделы ИИ-помощника (39 направлений)", 2)
    topics = sorted(SITE_TOPICS.values(), key=lambda x: x["title"])
    for t in topics:
        bullet(doc, f"{t['title']} — {t['url']}")

    heading(doc, "Полный перечень вопросов «Вопрос-ответ» (48)", 2)
    para(doc, "На сайте — развёрнутые ответы с правовым обоснованием; в боте — краткая выжимка.", bold=False)
    for i, item in enumerate(VOPROSY_ITEMS, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(item["label"])


def part_ops(doc: Document) -> None:
    heading(doc, "Часть IV. Приёмка, IT и инструменты куратора", 1)

    heading(doc, "Чек-лист регрессии (Telegram, ~3 мин)", 2)
    checks = [
        ("Филина", "Справочник, не «филиалы»"),
        ("телефон Миронова", "Справочник"),
        ("где устав / устав", "Документы СРО на сайте"),
        ("размеры взносов", "Взносы, не партнёр НОСО"),
        ("база законов, техрегулирование", "Разделы сайта, не справочник"),
        ("еврокоды что это", "FAQ про Еврокоды"),
        ("Тестовый ИНН", "Карточка реестра"),
    ]
    add_table(doc, ["Запрос", "Ожидание"], checks)

    heading(doc, "На ПК перед выкладкой", 2)
    bullet(doc, "python routing_regression.py — автопроверка маршрутизации.")
    bullet(doc, "Передача IT: IT_ZAYAVKA_DEPLOY.md, без config_keys.py в почте.")

    heading(doc, "Юридическая оговорка", 2)
    para(
        doc,
        "Бот публикует ссылки и материалы с официального сайта СРО и партнёрских реестров. "
        "Окончательная трактовка норм и индивидуальные случаи — только через специалистов "
        "Ассоциации и документы на https://www.srogen.ru/",
    )


def part_regiony(doc: Document) -> None:
    heading(doc, "Часть V. Зачем бот регионам и другим часовым поясам", 1)

    para(
        doc,
        "Аппарат Ассоциации ориентирован на московский рабочий день (UTC+3). Члены и "
        "кандидаты во Владивостоке (+7 ч), Хабаровске, Якутске (+6), Красноярске (+4) "
        "не могут звонить «когда удобно Москве», не сдвигая свой график на вечер или ночь.",
    )
    add_table(
        doc,
        ["Город (пример)", "Когда в Москве 10:00–18:00"],
        [
            ["Владивосток / Хабаровск", "17:00–01:00 (местное)"],
            ["Якутск", "16:00–00:00"],
            ["Красноярск", "14:00–22:00"],
            ["Екатеринбург", "12:00–20:00"],
        ],
    )
    para(
        doc,
        "Telegram-бот доступен круглосуточно: вступающая организация вечером по местному "
        "времени получает «как вступить», «размеры взносов», «какие документы» и ссылки "
        "на srogen.ru без ожидания линии. Действующий член — месяц проверки по ИНН, "
        "бланки, устав, НОК/НРС. Звонок в Ассоциацию остаётся для нестандартных кейсов, "
        "но клиент уже подготовлен.",
    )

    heading(doc, "Сценарии", 2)
    add_table(
        doc,
        ["Ситуация", "Что даёт бот"],
        [
            ["Новая компания хочет вступить", "FAQ + ИИ + ссылки на документы и взносы в любое время"],
            ["«Когда проверка?»", "Поиск по ИНН → план проверок"],
            ["Нужен бланк", "7 форм в чате или раздел FAQ"],
            ["Другой часовой пояс", "Не нужно совпадать с 10–18 МСК"],
        ],
    )

    quote = doc.add_paragraph()
    quote.paragraph_format.left_indent = Cm(1)
    quote.add_run(
        "«Член или компания, которая хочет вступить в СРО «ГЕН», не должна подстраиваться "
        "под московский часовой пояс. Бот даёт в любое время ответ на типовой вопрос и "
        "ссылку на официальный документ.»"
    ).italic = True

    para(
        doc,
        "Для стабильной работы ночью во Владивостоке нужен VPS 24/7 (см. часть IV и "
        "IT_ZAYAVKA_DEPLOY.md), а не только офисный ПК.",
    )


def part_en(doc: Document) -> None:
    heading(doc, "Часть VI. Перспектива: англоязычная версия (предложение)", 1)

    para(
        doc,
        "Статус: не начинать без решения руководства. Русскоязычный бот — основной и "
        "юридически значимый; английский — навигация для партнёров и не русскоязычных "
        "пользователей с отсылкой на srogen.ru (RU).",
    )

    add_table(
        doc,
        ["Вариант", "Срок", "Содержание"],
        [
            ["A — мини-EN (рекомендуется для пилота)", "1–2 нед.", "8–10 пунктов меню + ссылки + disclaimer"],
            ["B — рабочий EN", "1–2 мес.", "A + краткие тексты: вступление, взносы, NOK/NRS"],
            ["C — полный паритет", "долго", "Перевод 48 FAQ + юридическая вычитка"],
        ],
    )

    heading(doc, "Дисклеймер для EN-бота (текст)", 2)
    box = doc.add_paragraph()
    box.paragraph_format.left_indent = Cm(1)
    box.add_run(
        "This bot provides general guidance and links to the official website of "
        "SRO Association «GEN» (srogen.ru). Legal documents and binding information "
        "are published in Russian on the official website. For individual cases, "
        "please contact the Association: +7 (495) 775-81-11, info@srogen.ru."
    ).italic = True

    heading(doc, "Роли и запуск", 2)
    bullet(doc, "Руководство — утвердить A или B; второй бот в BotFather (рекомендуется).")
    bullet(doc, "Куратор — тексты EN, приёмка; IT — VPS по IT_ZAYAVKA_DEPLOY.md.")
    bullet(doc, "После стабильного RU на VPS — пилот EN.")

    para(
        doc,
        "Подробное ТЗ также в файле OTCHET_TZ_ANGLIYSKIY_BOT.docx (можно не печатать отдельно — "
        "суть включена в настоящий отчёт).",
    )


def main() -> None:
    doc = Document()
    set_doc_defaults(doc)
    add_title_page(doc)
    part_curator(doc)
    doc.add_page_break()
    part_bot(doc)
    doc.add_page_break()
    part_baza(doc)
    doc.add_page_break()
    part_ops(doc)
    doc.add_page_break()
    part_regiony(doc)
    doc.add_page_break()
    part_en(doc)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(
        f"Контакты Ассоциации: +7 (495) 775-81-11 · info@srogen.ru · {BOT_VERSION}"
    ).font.size = Pt(10)

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
