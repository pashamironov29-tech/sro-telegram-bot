"""ИИ-помощник контролёра: голос, разбор загруженных документов, выжимка в чат."""

from __future__ import annotations

import base64
import io
import logging
import re
import tempfile
import time
from pathlib import Path

import requests

try:
    from config_keys import GROQ_API_KEY as _GROQ
except Exception:
    _GROQ = ""

try:
    from config_keys import OPENROUTER_API_KEY as _OR_KEY
    from config_keys import OPENROUTER_MODEL as _OR_MODEL
except Exception:
    _OR_KEY = ""
    _OR_MODEL = "openai/gpt-4.1-mini"

try:
    from config_keys import OPENROUTER_DOC_MODEL as _OR_DOC_MODEL
except Exception:
    _OR_DOC_MODEL = ""

try:
    from bot_disclaimers import OFFICIAL_SOURCE_DISCLAIMER
except Exception:
    OFFICIAL_SOURCE_DISCLAIMER = (
        "ℹ️ Ориентир по тексту · не замена официальному документу и юристу."
    )

CONTROLLER_AI_BUTTON = "🎙 ИИ-помощник"
CONTROLLER_AI_HINT = (
    "🎙 <b>ИИ-помощник контролёра</b>\n\n"
    "Отправьте сюда:\n"
    "• <b>голосовое</b> — распознаю речь и отвечу\n"
    "• <b>файл</b> (PDF, Word .doc/.docx, TXT) — вытащу ключевое в чат\n"
    "• <b>фото</b> документа — постараюсь прочитать и кратко изложить\n"
    "• <b>текст</b> — вопрос по проверке / по уже присланному файлу\n\n"
    "После файла можно спросить: «какие сроки?», «кто подписал?» — "
    "отвечу по этому документу.\n\n"
    f"{OFFICIAL_SOURCE_DISCLAIMER}\n\n"
    "Выход — «⬅️ Назад в меню»."
)

_controller_ai_mode: set[int] = set()
# Последний разобранный текст документа (для уточняющих вопросов)
_last_upload_text: dict[int, str] = {}
_last_upload_name: dict[int, str] = {}

_MAX_DOC_CHARS = 60000
_MAX_REPLY_CHARS = 3500
def _openrouter_base() -> str:
    try:
        from config_keys import OPENROUTER_BASE as _b
    except Exception:
        _b = ""
    b = (_b or "https://openrouter.ai/api/v1").strip().rstrip("/")
    return b or "https://openrouter.ai/api/v1"


OPENROUTER_URL = _openrouter_base() + "/chat/completions"
OPENROUTER_DEFAULT_MODEL = "openai/gpt-4.1-mini"
OPENROUTER_DEFAULT_DOC_MODEL = "google/gemini-2.5-flash"
GROQ_CHAT_URL = "https://api.groq.com/openai/v1/chat/completions"


def _doc_model() -> str:
    """Сканы/почерк: Gemini Flash; запас — GPT-4.1. Не mini из FAQ."""
    for cand in (_OR_DOC_MODEL, OPENROUTER_DEFAULT_DOC_MODEL, _OR_MODEL, OPENROUTER_DEFAULT_MODEL):
        s = (cand or "").strip()
        if s:
            return s
    return OPENROUTER_DEFAULT_DOC_MODEL


def is_controller_ai_mode(chat_id: int) -> bool:
    try:
        return int(chat_id) in _controller_ai_mode
    except (TypeError, ValueError):
        return False


def enter_controller_ai_mode(chat_id: int) -> None:
    try:
        _controller_ai_mode.add(int(chat_id))
    except (TypeError, ValueError):
        pass


def exit_controller_ai_mode(chat_id: int) -> None:
    try:
        cid = int(chat_id)
    except (TypeError, ValueError):
        return
    _controller_ai_mode.discard(cid)


def fix_known_sro_ocr_names(text: str) -> str:
    """OCR часто читает Д как В в инициалах известных сотрудников СРО."""
    if not text:
        return text
    out = text
    # Дмитрий Владимирович Глебов — председатель КК / руководитель аппарата
    out = re.sub(
        r"\b[ВB]\.\s*[ВB]\.\s*(Глебов\w*)",
        r"Д.В. \1",
        out,
        flags=re.IGNORECASE,
    )
    out = re.sub(
        r"\b[ВB][ВB]\.\s*(Глебов\w*)",
        r"Д.В. \1",
        out,
        flags=re.IGNORECASE,
    )
    return out


def remember_upload(chat_id: int, text: str, name: str = "документ") -> None:
    try:
        cid = int(chat_id)
    except (TypeError, ValueError):
        return
    cleaned = fix_known_sro_ocr_names((text or "").strip())
    if len(cleaned) > _MAX_DOC_CHARS:
        cleaned = cleaned[:_MAX_DOC_CHARS] + "\n…"
    _last_upload_text[cid] = cleaned
    _last_upload_name[cid] = (name or "документ").strip()[:120]


def get_upload_context(chat_id: int) -> tuple[str, str]:
    try:
        cid = int(chat_id)
    except (TypeError, ValueError):
        return "", ""
    return _last_upload_text.get(cid, ""), _last_upload_name.get(cid, "")


def _html_escape(s: str) -> str:
    return (
        (s or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _trim_reply(text: str) -> str:
    t = (text or "").strip()
    if len(t) <= _MAX_REPLY_CHARS:
        return t
    return t[: _MAX_REPLY_CHARS - 20] + "\n…"


def _chat_completion(
    messages: list[dict],
    max_tokens: int = 1200,
    *,
    model: str | None = None,
    plugins: list | None = None,
    timeout: int = 90,
) -> str:
    or_key = (_OR_KEY or "").strip()
    model = (model or _OR_MODEL or OPENROUTER_DEFAULT_MODEL).strip() or OPENROUTER_DEFAULT_MODEL
    if or_key:
        last_exc: BaseException | None = None
        for attempt in range(3):
            try:
                r = requests.post(
                    OPENROUTER_URL,
                    headers={
                        "Authorization": f"Bearer {or_key}",
                        "Content-Type": "application/json",
                        "HTTP-Referer": "https://www.srogen.ru",
                        "X-Title": "SRO GOLD Controller AI",
                    },
                    json={
                        "model": model,
                        "messages": messages,
                        "temperature": 0.1,
                        "max_tokens": max_tokens,
                        **({"plugins": plugins} if plugins else {}),
                    },
                    timeout=timeout,
                )
                r.raise_for_status()
                return r.json()["choices"][0]["message"]["content"].strip()
            except (requests.ConnectionError, requests.Timeout) as exc:
                last_exc = exc
                logging.warning("OpenRouter chat attempt %s failed: %s", attempt + 1, exc)
                if attempt < 2:
                    time.sleep(1.2 * (attempt + 1))
                    continue
                raise
        if last_exc:
            raise last_exc

    groq = (_GROQ or "").strip()
    if not groq:
        raise RuntimeError("no_llm_key")
    # Vision-сообщения на Groq не отправляем — только текст
    plain = []
    for m in messages:
        content = m.get("content")
        if isinstance(content, list):
            parts = []
            for part in content:
                if isinstance(part, dict) and part.get("type") == "text":
                    parts.append(part.get("text") or "")
            content = "\n".join(parts)
        plain.append({"role": m["role"], "content": content or ""})
    r = requests.post(
        GROQ_CHAT_URL,
        headers={
            "Authorization": f"Bearer {groq}",
            "Content-Type": "application/json",
        },
        json={
            "model": "llama-3.1-8b-instant",
            "messages": plain,
            "temperature": 0.1,
            "max_tokens": max_tokens,
        },
        timeout=60,
    )
    r.raise_for_status()
    return r.json()["choices"][0]["message"]["content"].strip()



def _audio_format(filename: str) -> str:
    name = (filename or "voice.ogg").lower()
    for ext in ("ogg", "mp3", "wav", "m4a", "webm", "flac", "aac"):
        if name.endswith("." + ext):
            return ext
    return "ogg"


def _transcribe_openrouter(data: bytes, filename: str) -> str:
    or_key = (_OR_KEY or "").strip()
    if not or_key:
        raise RuntimeError("no_openrouter_key")
    fmt = _audio_format(filename)
    model = "openai/whisper-1"
    b64 = base64.b64encode(data).decode("ascii")
    r = requests.post(
        _openrouter_base() + "/audio/transcriptions",
        headers={
            "Authorization": f"Bearer {or_key}",
            "HTTP-Referer": "https://www.srogen.ru",
            "X-Title": "SRO GOLD Controller STT",
            "Content-Type": "application/json",
        },
        json={
            "model": model,
            "language": "ru",
            "input_audio": {"data": b64, "format": fmt},
        },
        timeout=120,
    )
    if r.status_code >= 400:
        files = {"file": (filename or "voice.ogg", data, "audio/ogg")}
        r = requests.post(
            _openrouter_base() + "/audio/transcriptions",
            headers={
                "Authorization": f"Bearer {or_key}",
                "HTTP-Referer": "https://www.srogen.ru",
                "X-Title": "SRO GOLD Controller STT",
            },
            files=files,
            data={"model": model, "language": "ru"},
            timeout=120,
        )
    if r.status_code >= 400:
        raise RuntimeError(f"openrouter_stt_{r.status_code}")
    try:
        return (r.json().get("text") or "").strip()
    except Exception:
        return (r.text or "").strip()


_local_whisper_model = None


def _transcribe_local(data: bytes, filename: str) -> str:
    """Локальный Whisper на VPS — если OpenRouter режет по security policy."""
    global _local_whisper_model
    import tempfile
    from pathlib import Path as _P

    try:
        from faster_whisper import WhisperModel
    except Exception as exc:
        raise RuntimeError("no_local_whisper") from exc

    if _local_whisper_model is None:
        # tiny/base — хватает для диктовок контролёра на CPU
        _local_whisper_model = WhisperModel(
            "tiny",
            device="cpu",
            compute_type="int8",
        )

    suffix = "." + _audio_format(filename)
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(data)
        tmp_path = tmp.name
    try:
        segments, _info = _local_whisper_model.transcribe(
            tmp_path,
            language="ru",
            vad_filter=True,
        )
        parts = [s.text.strip() for s in segments if (s.text or "").strip()]
        return " ".join(parts).strip()
    finally:
        try:
            _P(tmp_path).unlink(missing_ok=True)
        except Exception:
            pass


def transcribe_voice(data: bytes, filename: str = "voice.ogg") -> str:
    """STT: OpenRouter Whisper, при 403/сбое — локальный Whisper на сервере."""
    if not data:
        raise RuntimeError("empty_audio")

    # 1) Платный OpenRouter (если ключ жив и не режется политикой)
    try:
        text = _transcribe_openrouter(data, filename or "voice.ogg")
        if text:
            return text
    except Exception as exc:
        logging.warning("OpenRouter STT unavailable: %s", exc)

    # 2) Локально на Москве (не зависит от блокировок OR)
    text = _transcribe_local(data, filename or "voice.ogg")
    return text




def extract_text_from_pdf_openrouter(data: bytes, filename: str = "document.pdf") -> str:
    """Весь PDF разом: OpenRouter file-parser (mistral-ocr) + модель для документов."""
    or_key = (_OR_KEY or "").strip()
    if not or_key or not data:
        return ""
    if len(data) > 12 * 1024 * 1024:
        return ""
    b64 = base64.b64encode(data).decode("ascii")
    data_url = f"data:application/pdf;base64,{b64}"
    prompt = (
        "Это PDF для контролёра СРО (часто скан трудовой книжки или акта). "
        "Спиши ВЕСЬ читаемый текст со всех страниц. "
        "Для трудовой: каждое место работы — организация, должность, "
        "даты приёма и увольнения, номер приказа, статья ТК — если видно. "
        "Не выдумывай неразборчивое: пиши «неразборчиво». По-русски."
    )
    payload = {
        "model": _doc_model(),
        "temperature": 0.1,
        "max_tokens": 4000,
        "plugins": [{"id": "file-parser", "pdf": {"engine": "mistral-ocr"}}],
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {
                        "type": "file",
                        "file": {
                            "filename": filename or "document.pdf",
                            "file_data": data_url,
                        },
                    },
                ],
            }
        ],
    }
    try:
        r = requests.post(
            OPENROUTER_URL,
            headers={
                "Authorization": f"Bearer {or_key}",
                "Content-Type": "application/json",
                "HTTP-Referer": "https://www.srogen.ru",
                "X-Title": "SRO GOLD Controller AI",
            },
            json=payload,
            timeout=180,
        )
        r.raise_for_status()
        body = r.json()
    except Exception as exc:
        logging.warning("OpenRouter PDF OCR failed: %s", exc)
        return ""
    if body.get("error"):
        logging.warning("OpenRouter PDF OCR error: %s", body.get("error"))
        return ""
    msg = ((body.get("choices") or [{}])[0].get("message") or {})
    chunks: list[str] = []
    for ann in msg.get("annotations") or []:
        file_info = (ann or {}).get("file") or {}
        for part in file_info.get("content") or []:
            if isinstance(part, dict) and part.get("type") == "text" and part.get("text"):
                chunks.append(str(part["text"]).strip())
    content = (msg.get("content") or "").strip()
    if content:
        chunks.append(content)
    text = "\n\n".join(x for x in chunks if x).strip()
    if len(text) < 80:
        return ""
    return text


def extract_text_from_pdf_scan(data: bytes, max_pages: int = 12) -> str:
    """Скан-PDF: все страницы одним запросом в vision (почерк). Не Mistral OCR."""
    try:
        import fitz  # PyMuPDF
    except Exception as exc:
        raise RuntimeError("no_pymupdf") from exc
    if not data:
        return ""
    doc = fitz.open(stream=data, filetype="pdf")
    images: list[str] = []
    try:
        total = len(doc)
        n = min(total, max_pages)
        for i in range(n):
            page = doc[i]
            pix = page.get_pixmap(matrix=fitz.Matrix(2.2, 2.2), alpha=False)
            png = pix.tobytes("png")
            images.append(base64.b64encode(png).decode("ascii"))
    finally:
        doc.close()
    if not images:
        return ""
    prompt = (
        "Это страницы документа СРО (заявление, инфолист, трудовая, письмо). "
        f"Страниц в запросе: {len(images)}. "
        "Спиши текст как на бланке, по буквам. Печатный текст не называй плохой копией. "
        "SCAN_QUALITY=LOW только если страница тёмная, боком или почерк почти не читается; "
        "иначе SCAN_QUALITY=OK. "
        "Фамилию и названия фирм не угадывай похожими словами. Спорное — два варианта. "
        "Инициалы: не путай Д и В (Д.В. Глебов ≠ В.В. Глебов). "
        "Если это трудовая: каждое место работы — организация, должность, даты, приказ. "
        "Неразборчивое помечай «неразборчиво». Не выдумывай записи. По-русски."
    )
    content: list[dict] = [{"type": "text", "text": prompt}]
    for b64 in images:
        content.append(
            {
                "type": "image_url",
                "image_url": {"url": f"data:image/png;base64,{b64}"},
            }
        )
    messages = [{"role": "user", "content": content}]
    last = ""
    for model in (_doc_model(), "openai/gpt-4.1"):
        try:
            last = _chat_completion(
                messages,
                max_tokens=4000,
                model=model,
                timeout=180,
            )
            if last and len(last.strip()) >= 80:
                return last.strip()
        except Exception as exc:
            logging.warning("pdf vision %s failed: %s", model, exc)
    return (last or "").strip()



def _soffice_bin() -> str | None:
    import shutil

    for name in ("soffice", "libreoffice"):
        found = shutil.which(name)
        if found:
            return found
    win = Path(r"C:\Program Files\LibreOffice\program\soffice.exe")
    if win.is_file():
        return str(win)
    return None


def _decode_loose_bytes(raw: bytes) -> str:
    for enc in ("utf-8", "utf-16", "cp1251", "latin-1"):
        try:
            text = raw.decode(enc)
            if text.strip():
                return text
        except Exception:
            continue
    return raw.decode("utf-8", errors="ignore")


def _doc_via_soffice(data: bytes) -> str:
    """Старый .doc → текст через LibreOffice, если он есть на машине."""
    import subprocess
    import tempfile

    bin_path = _soffice_bin()
    if not bin_path or not data:
        return ""
    try:
        with tempfile.TemporaryDirectory(prefix="sro_doc_") as tmp:
            tmp_path = Path(tmp)
            src = tmp_path / "upload.doc"
            src.write_bytes(data)
            profile = tmp_path / "lo_profile"
            profile.mkdir()
            cmd = [
                bin_path,
                "--headless",
                "--nologo",
                "--nofirststartwizard",
                "--norestore",
                f"-env:UserInstallation={profile.as_uri()}",
                "--convert-to",
                "txt:Text",
                "--outdir",
                str(tmp_path),
                str(src),
            ]
            subprocess.run(
                cmd,
                check=False,
                timeout=90,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
            produced = tmp_path / "upload.txt"
            if not produced.is_file():
                cands = list(tmp_path.glob("*.txt"))
                produced = cands[0] if cands else None
            if not produced or not produced.is_file():
                return ""
            return _decode_loose_bytes(produced.read_bytes()).strip()
    except Exception as exc:
        logging.warning("soffice .doc convert failed: %s", exc)
        return ""


def _doc_via_ole(data: bytes) -> str:
    """Грубый разбор WordDocument stream (бланки СРО на cp/utf-16)."""
    import struct

    try:
        import olefile
    except Exception:
        return ""
    if not data:
        return ""
    try:
        ole = olefile.OleFileIO(io.BytesIO(data))
        try:
            if not ole.exists("WordDocument"):
                return ""
            stream = ole.openstream("WordDocument").read()
        finally:
            ole.close()
    except Exception as exc:
        logging.warning("olefile .doc failed: %s", exc)
        return ""

    def keep(ch: int) -> bool:
        if 0x0400 <= ch <= 0x04FF:
            return True
        if ch in (0x09, 0x0A, 0x0D, 0x20, 0xA0):
            return True
        if 0x21 <= ch <= 0x7E:
            return True
        return ch in (
            0xAB, 0xBB, 0x2013, 0x2014, 0x201C, 0x201D,
            0x2026, 0x2116, 0x00B0, 0x00A7,
        )

    blocks: list[str] = []
    buf = bytearray()
    i = 0
    n = len(stream)

    def flush() -> None:
        nonlocal buf
        if len(buf) < 12:
            buf = bytearray()
            return
        try:
            chunk = buf.decode("utf-16le")
        except Exception:
            buf = bytearray()
            return
        if any("\u0400" <= c <= "\u04FF" for c in chunk):
            blocks.append(chunk)
        buf = bytearray()

    while i + 1 < n:
        ch = stream[i] | (stream[i + 1] << 8)
        if keep(ch):
            buf += struct.pack("<H", ch)
            i += 2
        else:
            flush()
            i += 1
    flush()
    text = "\n".join(blocks)
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text_from_doc(data: bytes) -> str:
    """Текст из старого Word .doc (бланки СРО)."""
    if not data:
        return ""
    if data[:2] == b"PK":
        return extract_text_from_bytes(data, "upload.docx")
    text = _doc_via_soffice(data)
    if len(text) >= 80:
        return text
    ole = _doc_via_ole(data)
    if len(ole) > len(text):
        return ole
    return (text or ole).strip()


def extract_text_from_bytes(data: bytes, filename: str = "") -> str:
    """Текст из PDF / DOC / DOCX / TXT."""
    name = (filename or "").lower()
    if not data:
        return ""

    if name.endswith(".txt") or name.endswith(".csv") or name.endswith(".md"):
        for enc in ("utf-8", "cp1251", "latin-1"):
            try:
                return data.decode(enc)
            except Exception:
                continue
        return data.decode("utf-8", errors="ignore")

    if name.endswith(".doc") and not name.endswith(".docx"):
        return extract_text_from_doc(data)

    if name.endswith(".docx"):
        from docx import Document

        doc = Document(io.BytesIO(data))
        parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    if name.endswith(".pdf"):
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        parts = []
        for page in reader.pages[:40]:
            try:
                t = page.extract_text() or ""
            except Exception:
                t = ""
            if t.strip():
                parts.append(t)
        text = "\n".join(parts).strip()
        # Мало текста или «кривой» слой сканера → OCR по картинкам страниц
        if len(text) >= 80 and not looks_like_garbled_pdf_text(text):
            return text
        try:
            scanned = extract_text_from_pdf_scan(data)
            if len(scanned) < 80:
                scanned = extract_text_from_pdf_openrouter(data, filename)

        except Exception as exc:
            logging.warning("pdf scan OCR failed: %s", exc)
            scanned = ""
        if scanned:
            if text and not looks_like_garbled_pdf_text(text):
                return (text + "\n\n" + scanned).strip()
            return scanned
        return text

    # По расширению не узнали — пробуем как текст
    sample = data[:200]
    if b"\\x00" not in sample:
        return data.decode("utf-8", errors="ignore")
    return ""



def looks_like_garbled_pdf_text(text: str) -> bool:
    """Встроенный слой сканера: латиница в русских словах, обрывки по буквам."""
    raw = text or ""
    if len(raw.strip()) < 80:
        return True
    tokens = re.findall(r"\S+", raw)
    if not tokens:
        return True

    def _letters(tok: str) -> str:
        return re.sub(r"[^\wА-Яа-яЁё]", "", tok, flags=re.UNICODE)

    short = sum(1 for t in tokens if len(_letters(t)) <= 1)
    if len(tokens) >= 60 and short / len(tokens) >= 0.35:
        return True
    mixed = re.findall(
        r"(?:[А-Яа-яЁё]{2,}[A-Za-z]+|[A-Za-z]{2,}[А-Яа-яЁё]+)",
        raw,
    )
    if len(mixed) >= 6:
        return True
    letters = [ch for ch in raw if ch.isalpha()]
    if len(letters) >= 200:
        cyr = sum(
            1 for ch in letters if "А" <= ch.upper() <= "Я" or ch.upper() == "Ё"
        )
        if cyr / len(letters) < 0.62:
            return True
    return False


def _years_in_text(text: str) -> list[int]:
    years: list[int] = []
    for m in re.finditer(r"\b(19|20)\d{2}\b", text or ""):
        try:
            years.append(int(m.group(0)))
        except ValueError:
            continue
    return years


def looks_like_readable_typed_doc(text: str) -> bool:
    """Печатное заявление / инфолист: ИНН есть, каши нет — не «плохая копия»."""
    raw = text or ""
    if looks_like_hallucinated_career(raw):
        return False
    if raw.lower().count("неразборчив") >= 3:
        return False
    has_inn = bool(re.search(r"\b\d{10}\b", raw))
    has_marker = bool(
        re.search(
            r"(ИНН|ОГРН|ООО|ИНФОРМАЦИОНН|заявлен)",
            raw,
            flags=re.I,
        )
    )
    return has_inn and has_marker


def looks_like_hallucinated_career(text: str) -> bool:
    """ИИ «рисует» карьеру: приказы подряд и даты далеко в будущем."""
    raw = text or ""
    if not raw.strip():
        return False
    years = _years_in_text(raw)
    if any(y >= 2035 for y in years):
        return True
    orders = len(re.findall(r"приказ\s*№", raw, flags=re.I))
    first_days = re.findall(r"\b01\.(\d{2})\.(\d{4})\b", raw)
    if orders >= 5 and len(first_days) >= 6:
        return True
    if len(first_days) >= 10:
        return True
    return False


def looks_like_unreliable_scan(text: str, *, filename: str = "") -> bool:
    """Плохая копия или недостоверный разбор — нельзя выдавать как факт."""
    raw = (text or "").lower()
    if not raw:
        return True
    if looks_like_hallucinated_career(text):
        return True
    if looks_like_readable_typed_doc(text):
        return False
    n_bad = raw.count("неразборчив")
    if n_bad >= 2:
        return True
    if "ооо «неразборчиво»" in raw or 'ооо "неразборчиво"' in raw:
        return True
    if "scan_quality=low" in raw and n_bad >= 1:
        return True
    fname = (filename or "").lower()
    if ("труд" in fname or fname in ("tk.pdf", "тк.pdf")) and n_bad >= 1:
        return True
    return False


_UNRELIABLE_SCAN_BANNER = (
    "⚠️ <b>Плохая копия / нечёткий скан.</b>\n\n"
    "В этом документе <b>всё правильно разобрать не получится</b>: "
    "тёмная ксерокопия, мелкий почерк, разворот боком.\n\n"
    "ИИ мог ошибиться в <b>ФИО</b>, <b>названиях организаций</b> и <b>датах</b> "
    "или придумать записи, которых в книжке нет.\n\n"
    "⛔ <b>Не используйте этот разбор как факт</b> — только сверка с оригиналом.\n"
    "Лучше пришлите <b>фото разворота при свете</b> (крупно, без блика).\n\n"
    "<i>Ориентир по тексту · сверь с оригиналом.</i>"
)


def _unreliable_scan_only() -> str:
    return _UNRELIABLE_SCAN_BANNER


def _strip_false_bad_copy_prefix(answer: str) -> str:
    body = (answer or "").strip()
    if not body:
        return body
    plain = re.sub(r"<[^>]+>", "", body)
    head = plain[:220].lower()
    if "плохая копия" not in head and "нечёткий скан" not in head:
        return body
    parts = re.split(r"\n\s*\n", body, maxsplit=1)
    if len(parts) == 2:
        return parts[1].strip()
    return body


def _with_scan_quality_banner(answer: str, source_text: str, *, filename: str = "") -> str:
    hall = looks_like_hallucinated_career(answer) or looks_like_hallucinated_career(
        source_text
    )
    if hall:
        return _unreliable_scan_only()
    if looks_like_readable_typed_doc(source_text) or looks_like_readable_typed_doc(
        answer
    ):
        return _strip_false_bad_copy_prefix(answer)
    src_bad = looks_like_unreliable_scan(source_text, filename=filename)
    ans_bad = looks_like_unreliable_scan(answer, filename=filename)
    if src_bad or ans_bad:
        return _unreliable_scan_only()
    body = (answer or "").strip()
    if "плохая копия" not in body.lower() and "нечёткий скан" not in body.lower():
        if (source_text or "").lower().count("неразборчив") >= 2:
            return _UNRELIABLE_SCAN_BANNER + "\n\n" + body
    return body


def summarize_upload_for_controller(text: str, filename: str = "документ") -> str:
    body = fix_known_sro_ocr_names((text or "").strip())
    if not body:
        return (
            "⚠️ Не удалось извлечь текст из файла. "
            "Пришлите PDF/Word с текстовым слоем или фото крупнее."
        )
    if len(body) > _MAX_DOC_CHARS:
        body = body[:_MAX_DOC_CHARS] + "\n…"

    system = (
        "Ты помощник контролёра СРО. По тексту загруженного документа "
        "выдели практичную информацию для проверки."
        " Пиши по-русски, маркированными блоками."
        " Не выдумывай факты, которых нет в тексте. Не угадывай даты и названия."
        " Печатное заявление, инфолист, письмо с ИНН/адресом — это нормальный документ: "
        "НЕ пиши «плохая копия» и не проси фото, даже если в тексте редкие опечатки OCR."
        " «Плохая копия» — только если в тексте много «неразборчиво» "
        "или это тёмная рукописная трудовая без названий организаций."
        " Если это читаемая трудовая — перечисли ВСЕ места работы по порядку: "
        "организация, должность, дата приёма, дата увольнения, основание. "
        "Не пиши «других сведений нет», если в тексте ещё есть записи."
        " Если это акт/письмо/выписка/инфолист — стороны, даты, ИНН, адреса, "
        "предмет, сроки, требования, подписи, специалисты НРС — что найдётся."
        " Инициалы Д и В не путай. Председатель КК — Д.В. Глебов, не В.В."
        " В конце одной строкой: «Ориентир по тексту · сверь с оригиналом.»"
    )
    if looks_like_unreliable_scan(body, filename=filename):
        return _unreliable_scan_only()

    user = f"Имя файла: {filename}\n\nТекст:\n{body}"
    try:
        answer = _chat_completion(
            [
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
            max_tokens=2200,
            model=_doc_model(),
            timeout=120,
        )
    except RuntimeError:
        return (
            "⚠️ Нет ключа ИИ (OPENROUTER_API_KEY или GROQ_API_KEY).\n"
            "Текст из файла получен, но разобрать его не удалось."
        )
    except Exception as exc:
        logging.warning("controller_ai summarize failed: %s", exc)
        # Fallback: первые абзацы без ИИ
        if looks_like_unreliable_scan(body, filename=filename):
            return _unreliable_scan_only()
        preview = _html_escape(_trim_reply(body[:1200]))
        return (
            f"📄 <b>{_html_escape(filename)}</b>\n\n"
            f"ИИ временно недоступен. Фрагмент текста (может быть неточным):\n\n{preview}"
        )
    return fix_known_sro_ocr_names(
        _with_scan_quality_banner(
            _trim_reply(_html_escape(answer)), body, filename=filename
        )
    )


def answer_about_upload(question: str, chat_id: int) -> str:
    doc_text, doc_name = get_upload_context(chat_id)
    if not doc_text:
        return (
            "Сначала пришлите <b>файл или фото</b> документа — "
            "потом можно задавать уточнения по нему."
        )
    system = (
        "Ты помощник контролёра СРО. Отвечай ТОЛЬКО по тексту присланного документа. "
        "Если спрашивают, где человек работал — перечисли все места из текста, не два первых. "
        "Не выдумывай даты в будущем и не подставляй чужие названия. "
        "Если документ — плохая копия / много «неразборчиво»: сначала предупреди, "
        "что правильно разобрать не получится, и не достраивай карьеру. "
        "Если в тексте нет ответа — так и скажи. По-русски. "
        "В конце: «Ориентир по тексту · сверь с оригиналом.»"
    )
    user = (
        f"Документ: {doc_name}\n\n"
        f"Вопрос: {question}\n\n"
        f"Текст документа:\n{doc_text}"
    )
    try:
        answer = _chat_completion(
            [
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
            model=_doc_model(),
            timeout=120,
        )
    except Exception as exc:
        logging.warning("controller_ai answer_about_upload failed: %s", exc)
        return "⚠️ Не удалось ответить по документу. Попробуйте ещё раз."
    return fix_known_sro_ocr_names(
        _with_scan_quality_banner(
            _trim_reply(_html_escape(answer)), doc_text, filename=doc_name
        )
    )


def extract_text_from_image(data: bytes, mime: str = "image/jpeg") -> str:
    """OCR/выжимка с фото через OpenRouter vision (если есть ключ)."""
    or_key = (_OR_KEY or "").strip()
    if not or_key:
        raise RuntimeError("no_openrouter_for_vision")
    if not data:
        return ""
    b64 = base64.b64encode(data).decode("ascii")
    model = _doc_model()
    # Сканы документов — GPT-4.1, не mini из FAQ-чата
    prompt = (
        "Это страница документа СРО (заявление, инфолист, трудовая, письмо). "
        "Спиши текст по буквам. Печатный текст не называй плохой копией. "
        "Фамилию и названия организаций не заменяй похожим словом. Спорное — два варианта. "
        "Инициалы Д и В не путай (Д.В. Глебов, не В.В.). "
        "Если трудовая: организация, должность, даты, приказ. "
        "Неразборчивое — «неразборчиво». Не выдумывай. По-русски."
    )
    messages = [
        {
            "role": "user",
            "content": [
                {"type": "text", "text": prompt},
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:{mime};base64,{b64}"},
                },
            ],
        }
    ]
    return _chat_completion(messages, max_tokens=2500, model=model, timeout=120)




def controller_free_chat(question: str) -> str:
    """Свободный ответ для режима контролёра без подсказок разделов сайта."""
    q = (question or "").strip()
    if not q:
        return "Пустой запрос."

    ql = q.lower()
    # явный тест микрофона / распознавания — без ухода в планы проверок
    test_markers = (
        "1,2,3", "1 2 3", "раз два три", "проверка микрофон",
        "тест микрофон", "тест голос", "проверка связи", "просто тест",
    )
    if any(m in ql for m in test_markers) or re.fullmatch(r"(проверка|тест)([\s,\.\-]*\d+)*", ql):
        return (
            "✅ Голос/текст принят.\n"
            f"Распознано: <b>{_html_escape(q)}</b>\n\n"
            "Это похоже на проверку связи — подсказки по разделам сайта не даю. "
            "Задайте рабочий вопрос или пришлите документ."
        )

    system = (
        "Ты помощник контролёра СРО в Telegram. "
        "Отвечай кратко по-русски по сути вопроса. "
        "НЕ предлагай сам разделы сайта, планы проверок, бланки и «полезные ссылки», "
        "если пользователь явно об этом не спросил. "
        "Не выдавай юридических заключений. "
        "Если вопрос неясный или это тест — коротко уточни. "
        "В конце одной строкой: «Ориентир · сверь с официальным источником.»"
    )
    try:
        answer = _chat_completion(
            [
                {"role": "system", "content": system},
                {"role": "user", "content": q},
            ],
            max_tokens=900,
        )
    except RuntimeError:
        return (
            "⚠️ Нет доступа к ИИ (OpenROUTER). "
            "Проверьте OPENROUTER_BASE/ключ или напишите вопрос иначе."
        )
    except Exception as exc:
        logging.warning("controller_free_chat failed: %s", exc)
        return "⚠️ ИИ временно недоступен. Попробуйте ещё раз."
    return _trim_reply(_html_escape(answer))

def looks_like_doc_followup(question: str) -> bool:
    q = (question or "").lower()
    keys = (
        "в документе",
        "по документу",
        "из файла",
        "в файле",
        "в акте",
        "в письме",
        "кто подпис",
        "какой срок",
        "какие срок",
        "что написа",
        "найди в",
        "выпиши",
        "сумма",
        "инн",
    )
    return any(k in q for k in keys)


def safe_filename(name: str) -> str:
    name = (name or "document").strip() or "document"
    name = re.sub(r"[\\\\/:*?\"<>|]+", "_", name)
    return name[:180]
