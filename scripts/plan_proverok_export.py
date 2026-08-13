# -*- coding: utf-8 -*-
"""
Сборка планов проверок Word из «ИТОГИ КОНТРОЛЯ 2026.xlsx».

СРО определяется по цвету строки (легенда в строке 1).
Формат как в официальных plan_proverok_*.doc: УТВЕРЖДЕНО + таблица с блоками месяцев.

Запуск:
  py scripts/plan_proverok_export.py
  py scripts/plan_proverok_export.py --xlsx "C:\\path\\file.xlsx" --out "C:\\path\\out"
"""
from __future__ import annotations

import argparse
import re
import sys
import zipfile
from collections import defaultdict
from pathlib import Path

_ROOT = Path(__file__).resolve().parents[1]
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))

import openpyxl
from openpyxl.xml.functions import fromstring
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

THEME_FALLBACK = {
    0: "000000",
    1: "FFFFFF",
    2: "1F497D",
    3: "EEECE1",
    4: "4F81BD",
    5: "C0504D",
    6: "9BBB59",
    7: "8064A2",
    8: "4BACC6",
    9: "F79646",
}

MONTH_ORDER = [
    "ЯНВАРЬ",
    "ФЕВРАЛЬ",
    "МАРТ",
    "АПРЕЛЬ",
    "МАЙ",
    "ИЮНЬ",
    "ИЮЛЬ",
    "АВГУСТ",
    "СЕНТЯБРЬ",
    "ОКТЯБРЬ",
    "НОЯБРЬ",
    "ДЕКАБРЬ",
]

MONTH_MAP = {
    "ЯНВ": "ЯНВАРЬ",
    "ЯНВАРЬ": "ЯНВАРЬ",
    "ФЕВ": "ФЕВРАЛЬ",
    "ФЕВРАЛЬ": "ФЕВРАЛЬ",
    "МАРТ": "МАРТ",
    "АПР": "АПРЕЛЬ",
    "АПРЕЛЬ": "АПРЕЛЬ",
    "МАЙ": "МАЙ",
    "ИЮН": "ИЮНЬ",
    "ИЮНЬ": "ИЮНЬ",
    "ИЮЛ": "ИЮЛЬ",
    "ИЮЛЬ": "ИЮЛЬ",
    "АВГ": "АВГУСТ",
    "АВГУСТ": "АВГУСТ",
    "СЕН": "СЕНТЯБРЬ",
    "СЕНТ": "СЕНТЯБРЬ",
    "СЕНТЯБРЬ": "СЕНТЯБРЬ",
    "ОКТ": "ОКТЯБРЬ",
    "ОКТЯБРЬ": "ОКТЯБРЬ",
    "НОЯ": "НОЯБРЬ",
    "НОЯБРЬ": "НОЯБРЬ",
    "ДЕК": "ДЕКАБРЬ",
    "ДЕКА": "ДЕКАБРЬ",
    "ДЕКАБРЬ": "ДЕКАБРЬ",
}

MONTH_FILE = {
    "ЯНВАРЬ": "01_январь",
    "ФЕВРАЛЬ": "02_февраль",
    "МАРТ": "03_март",
    "АПРЕЛЬ": "04_апрель",
    "МАЙ": "05_май",
    "ИЮНЬ": "06_июнь",
    "ИЮЛЬ": "07_июль",
    "АВГУСТ": "08_август",
    "СЕНТЯБРЬ": "09_сентябрь",
    "ОКТЯБРЬ": "10_октябрь",
    "НОЯБРЬ": "11_ноябрь",
    "ДЕКАБРЬ": "12_декабрь",
}

# Метаданные шапки из официальных планов 2026 (пример в zip)
SRO_META = {
    "ОГПС": {
        "folder": "ОГПС",
        "decision": "решением Правления саморегулируемой организации Ассоциации «Объединение генеральных подрядчиков в строительстве»",
        "protocol": "Протокол №282  от «25» ноября 2025 года",
        "title_org": "Ассоциации «Объединение генеральных подрядчиков в строительстве»",
    },
    "ОГПП": {
        "folder": "ОГПП",
        "decision": "решением Правления саморегулируемой организации Ассоциации «Объединение градостроительного планирования и проектирования»",
        "protocol": "Протокол № 265 от «25» ноября 2025 года",
        "title_org": "Ассоциации «Объединение градостроительного планирования и проектирования»",
    },
    "ОСО": {
        "folder": "ОСО",
        "decision": "решением Правления саморегулируемой организации Ассоциации «Объединение строительных организаций среднего и малого бизнеса»",
        "protocol": "Протокол № 296 от «25» ноября 2025 года",
        "title_org": "Ассоциации «Объединение строительных организаций среднего и малого бизнеса»",
    },
    "ОГПО": {
        "folder": "ОГПО",
        "decision": "решением Правления саморегулируемой организации Ассоциации «Объединение градостроительных проектных организаций»",
        "protocol": "Протокол № 101 от «25» ноября 2025 года",
        "title_org": "Ассоциации «Объединение градостроительных проектных организаций»",
    },
    "ГеоИндустрия": {
        "folder": "Геоиндустрия",
        "decision": "решением Правления саморегулируемой организации Ассоциации «Объединение изыскателей «ГеоИндустрия»",
        "protocol": "Протокол № 120 от «25» ноября 2025 года",
        "title_org": "Ассоциации «Объединение изыскателей «ГеоИндустрия»",
    },
    "МГЕО": {
        "folder": "МГЕО",
        "decision": "решением Правления саморегулируемой организации Ассоциации «Межрегиональное объединение изыскателей «ГЕО»",
        "protocol": "Протокол № 100 от «25» ноября 2025 года",
        "title_org": "Ассоциации «Межрегиональное объединение изыскателей «ГЕО»",
    },
    "СПРОФ": {
        "folder": "СПРОФ",
        "decision": "решением Совета саморегулируемой организации Ассоциации «Содружество профессиональных проектировщиков в строительстве»",
        "protocol": "Протокол № 84 от «25» ноября 2025 года",
        "title_org": "Ассоциации «Содружество профессиональных проектировщиков в строительстве»",
    },
    "ПРИИС": {
        "folder": "ПРИИС",
        "decision": "решением Совета саморегулируемой организации Ассоциации «Профессионалы рынка инженерных изысканий в области строительства»",
        "protocol": "Протокол № 86 от «25» ноября 2025 года",
        "title_org": "Ассоциации «Профессионалы рынка инженерных изысканий в области строительства»",
    },
    "ОПП": {
        "folder": "ОПП",
        "decision": "решением Правления саморегулируемой организации Ассоциация организаций профессионального проектирования",
        "protocol": "Протокол №11/2025 от «25» ноября 2025 года",
        "title_org": "Ассоциация организаций профессионального проектирования",
    },
}


def load_theme_map(xlsx_path: Path) -> dict[int, str]:
    theme = dict(THEME_FALLBACK)
    with zipfile.ZipFile(xlsx_path) as z:
        names = [n for n in z.namelist() if "theme" in n.lower() and n.endswith(".xml")]
        if not names:
            return theme
        root = fromstring(z.read(names[0]))
        ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
        scheme = root.find(".//a:clrScheme", ns)
        if scheme is None:
            return theme
        order = [
            "dk1",
            "lt1",
            "dk2",
            "lt2",
            "accent1",
            "accent2",
            "accent3",
            "accent4",
            "accent5",
            "accent6",
        ]
        for i, tag in enumerate(order):
            el = scheme.find(f"a:{tag}", ns)
            if el is None:
                continue
            srgb = el.find(".//a:srgbClr", ns)
            sys = el.find(".//a:sysClr", ns)
            if srgb is not None:
                theme[i] = srgb.get("val")
            elif sys is not None:
                theme[i] = sys.get("lastClr") or theme[i]
    return theme


def apply_tint(rgb_hex: str, tint: float) -> str:
    if not tint:
        return rgb_hex.upper()
    r = int(rgb_hex[0:2], 16) / 255.0
    g = int(rgb_hex[2:4], 16) / 255.0
    b = int(rgb_hex[4:6], 16) / 255.0

    def one(c: float) -> float:
        if tint < 0:
            return c * (1.0 + tint)
        return c * (1.0 - tint) + tint

    rr = int(max(0.0, min(1.0, one(r))) * 255)
    gg = int(max(0.0, min(1.0, one(g))) * 255)
    bb = int(max(0.0, min(1.0, one(b))) * 255)
    return f"{rr:02X}{gg:02X}{bb:02X}"


def cell_color_key(cell, theme_map: dict[int, str]) -> str | None:
    fill = cell.fill
    if not fill or fill.patternType in (None, "none"):
        return None
    fg = fill.fgColor
    if fg is None:
        return None
    if fg.type == "rgb" and fg.rgb and fg.rgb not in ("00000000", "None"):
        rgb = fg.rgb
        if len(rgb) == 8:
            rgb = rgb[2:]
        return rgb.upper()
    if fg.type == "theme" and fg.theme is not None:
        base = theme_map.get(fg.theme, THEME_FALLBACK.get(fg.theme, "000000"))
        tint = getattr(fg, "tint", 0) or 0
        return apply_tint(base, tint)
    return None


def normalize_month(raw) -> str | None:
    if raw is None:
        return None
    s = str(raw).strip().upper().replace(".", "")
    if not s:
        return None
    if s in MONTH_MAP:
        return MONTH_MAP[s]
    for key, val in MONTH_MAP.items():
        if s.startswith(key):
            return val
    return None


def s(val) -> str:
    if val is None:
        return ""
    return str(val).strip()


def set_run_font(run, size=12, bold=False, underline=False):
    run.bold = bold
    run.underline = underline
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")


def set_cell_border(cell, sz="4", color="000000"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # remove old borders if any
    for child in list(tcPr):
        if child.tag == qn("w:tcBorders"):
            tcPr.remove(child)
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_shading(cell, fill_hex: str | None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:shd"):
            tcPr.remove(child)
    if not fill_hex:
        return
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def set_cell_width(cell, dxa: int):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:tcW"):
            tcPr.remove(child)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(dxa))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def set_cell_vertical_center(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:vAlign"):
            tcPr.remove(child)
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    tcPr.append(vAlign)


PLAN_ROW_NUMBERING_ABSTRACT_ID = 9901
PLAN_ROW_NUMBERING_NUM_ID = 9901


def setup_plan_row_numbering(doc: Document) -> int:
    """Автонумерация 1. 2. 3. — в Word пересчитывается при удалении строк."""
    numbering = doc.part.numbering_part.element
    for abstract in numbering.findall(qn("w:abstractNum")):
        if abstract.get(qn("w:abstractNumId")) == str(PLAN_ROW_NUMBERING_ABSTRACT_ID):
            return PLAN_ROW_NUMBERING_NUM_ID

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(PLAN_ROW_NUMBERING_ABSTRACT_ID))

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)

    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(lvl_text)

    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)

    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "0")
    ind.set(qn("w:hanging"), "0")
    p_pr.append(ind)
    lvl.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_pr.append(r_fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    r_pr.append(sz)
    lvl.append(r_pr)

    abstract_num.append(lvl)
    numbering.append(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(PLAN_ROW_NUMBERING_NUM_ID))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(PLAN_ROW_NUMBERING_ABSTRACT_ID))
    num.append(abstract_num_id)
    numbering.append(num)
    return PLAN_ROW_NUMBERING_NUM_ID


def _set_paragraph_numbering(paragraph, num_id: int, level: int = 0) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.insert(0, num_pr)


def write_auto_number_cell(cell, num_id: int, *, bold=False, size=12) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    _set_paragraph_numbering(p, num_id)
    run = p.add_run()
    set_run_font(run, size=size, bold=bold)
    set_cell_border(cell)
    set_cell_vertical_center(cell)


def write_cell_paragraphs(
    cell,
    lines: list[str],
    *,
    bold=False,
    size=12,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    underline=False,
):
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = align
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        run = p.add_run(line)
        set_run_font(run, size=size, bold=bold, underline=underline)
    set_cell_border(cell)
    set_cell_vertical_center(cell)


def merge_row(
    table,
    row_idx: int,
    text: str,
    *,
    bold=True,
    size=12,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    underline=False,
):
    row = table.rows[row_idx]
    cell0 = row.cells[0]
    for c in range(1, len(row.cells)):
        cell0.merge(row.cells[c])
    write_cell_paragraphs(
        cell0,
        [text],
        bold=bold,
        size=size,
        align=align,
        underline=underline,
    )
    set_cell_shading(cell0, None)
    return cell0


def write_data_row(
    table,
    row_idx: int,
    values: list[str],
    *,
    bold=False,
    size=12,
    center_cols=None,
    row_numbering_id: int | None = None,
):
    center_cols = center_cols or set()
    row = table.rows[row_idx]
    start_col = 0
    if row_numbering_id is not None:
        write_auto_number_cell(row.cells[0], row_numbering_id, bold=bold, size=size)
        start_col = 1
    for i, val in enumerate(values[start_col:], start=start_col):
        align = WD_ALIGN_PARAGRAPH.CENTER if i in center_cols else WD_ALIGN_PARAGRAPH.LEFT
        write_cell_paragraphs(
            row.cells[i],
            [val],
            bold=bold,
            size=size,
            align=align,
        )
        set_cell_shading(row.cells[i], None)


def set_document_defaults(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")


# Ширины колонок — как в эталоне
# Планы_проверок_2026_Word/СПРОФ/plan_proverok_СПРОФ_2026.docx
COL_WIDTHS_DXA = (682, 1583, 4485, 1755, 6638)  # sum 15143


def set_table_grid(table, widths_dxa: tuple[int, ...]):
    """Зафиксировать ширины столбцов через tblGrid (иначе Word «плывёт»)."""
    tbl = table._tbl
    tblGrid = tbl.tblGrid
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl.insert(0, tblGrid)
    for child in list(tblGrid):
        tblGrid.remove(child)
    for w in widths_dxa:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(w))
        tblGrid.append(gridCol)
    # ширина всей таблицы
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for child in list(tblPr):
        if child.tag == qn("w:tblW"):
            tblPr.remove(child)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(sum(widths_dxa)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def build_document(meta: dict, year: int, rows_by_month: dict[str, list[dict]], months_filter: list[str] | None = None) -> Document:
    doc = Document()
    set_document_defaults(doc)

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.5)

    # Шапка «УТВЕРЖДЕНО» справа (как в официальных планах)
    for text, bold in (
        ("УТВЕРЖДЕНО", True),
        (meta["decision"], False),
        (meta["protocol"], False),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.left_indent = Cm(14)
        run = p.add_run(text)
        set_run_font(run, size=12, bold=bold)

    # пустая строка
    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(0)
    gap.paragraph_format.space_after = Pt(6)

    if months_filter and len(months_filter) == 1:
        title2 = f"{meta['title_org']} на {months_filter[0].lower()} {year} года."
    else:
        title2 = f"{meta['title_org']} на {year} календарный год."

    # Заголовок ПЛАНА — над таблицей (как в эталоне СПРОФ)
    for text in (
        "План проверок членов саморегулируемой организации",
        title2,
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        run = p.add_run(text)
        set_run_font(run, size=12, bold=True)

    months = months_filter or [m for m in MONTH_ORDER if rows_by_month.get(m)]
    n_data = sum(len(rows_by_month.get(m, [])) for m in months)
    # 1 шапка + месяцы + данные (без строк заголовка внутри таблицы)
    n_rows = 1 + len(months) + n_data
    table = doc.add_table(rows=max(n_rows, 2), cols=5)
    table.autofit = False
    table.allow_autofit = False
    set_table_grid(table, COL_WIDTHS_DXA)

    for i, w in enumerate(COL_WIDTHS_DXA):
        set_cell_width(table.rows[0].cells[i], w)

    # Шапка таблицы
    headers = [
        ("№\nп/п", True),
        ("ОПФ", False),
        ("Наименование организации", False),
        ("ИНН", False),
        ("Адрес", False),
    ]
    for i, (label, twoline) in enumerate(headers):
        cell = table.rows[0].cells[i]
        lines = label.split("\n") if twoline else [label]
        write_cell_paragraphs(
            cell,
            lines,
            bold=True,
            size=12,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        set_cell_width(cell, COL_WIDTHS_DXA[i])
        set_cell_shading(cell, None)

    row_numbering_id = setup_plan_row_numbering(doc)
    ridx = 1
    for month in months:
        items = rows_by_month.get(month) or []
        if not items:
            continue
        merge_row(
            table,
            ridx,
            month,
            bold=True,
            size=12,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        set_cell_width(table.rows[ridx].cells[0], sum(COL_WIDTHS_DXA))
        ridx += 1
        for item in items:
            write_data_row(
                table,
                ridx,
                [
                    "",
                    item["opf"],
                    item["name"],
                    item["inn"],
                    item["address"],
                ],
                bold=False,
                size=12,
                center_cols=set(),
                row_numbering_id=row_numbering_id,
            )
            for i, w in enumerate(COL_WIDTHS_DXA):
                set_cell_width(table.rows[ridx].cells[i], w)
            ridx += 1

    tbl = table._tbl
    while len(table.rows) > ridx:
        tbl.remove(table.rows[-1]._tr)

    return doc


def read_excel(xlsx: Path):
    theme = load_theme_map(xlsx)
    wb = openpyxl.load_workbook(xlsx, data_only=False)
    wb_d = openpyxl.load_workbook(xlsx, data_only=True)

    legend = {}
    ws0 = wb[wb.sheetnames[0]]
    # Prefer ПЛАН sheet for legend
    legend_sheet = wb["ПЛАН"] if "ПЛАН" in wb.sheetnames else ws0
    for c in range(1, 20):
        name = legend_sheet.cell(1, c).value
        if not name or not isinstance(name, str):
            continue
        key = cell_color_key(legend_sheet.cell(1, c), theme)
        legend[key] = name.strip()

    # data[sro][month] = list of dicts
    data = defaultdict(lambda: defaultdict(list))
    seen = set()

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        ws_d = wb_d[sheet_name]
        for r in range(5, (ws.max_row or 0) + 1):
            name = s(ws_d.cell(r, 4).value)
            if not name:
                continue
            month = normalize_month(ws_d.cell(r, 1).value)
            if not month:
                continue
            key = cell_color_key(ws.cell(r, 4), theme) or cell_color_key(ws.cell(r, 3), theme)
            sro = legend.get(key)
            if sro is None and key is None:
                sro = legend.get(None)
            if not sro or sro not in SRO_META:
                continue
            inn = s(ws_d.cell(r, 5).value)
            arch = s(ws_d.cell(r, 7).value)
            dedupe = (sro, month, inn, arch, name, sheet_name)
            if dedupe in seen:
                continue
            seen.add(dedupe)
            address = s(ws_d.cell(r, 6).value) or s(ws_d.cell(r, 8).value)
            data[sro][month].append(
                {
                    "opf": s(ws_d.cell(r, 3).value),
                    "name": name,
                    "inn": inn,
                    "address": address,
                    "sheet": sheet_name,
                }
            )
    return data, legend


def safe_name(name: str) -> str:
    return re.sub(r'[<>:"/\\|?*]+', "_", name).strip()


def export(xlsx: Path, out_dir: Path, year: int = 2026, only_sro: str | None = None):
    out_dir.mkdir(parents=True, exist_ok=True)
    data, legend = read_excel(xlsx)
    summary = []

    for sro, months in sorted(data.items(), key=lambda x: x[0]):
        if only_sro:
            key = only_sro.strip().lower()
            folder = SRO_META.get(sro, {}).get("folder", "").lower()
            if key not in {sro.lower(), folder, safe_name(sro).lower()}:
                continue
        meta = SRO_META[sro]
        folder = out_dir / meta["folder"]
        folder.mkdir(parents=True, exist_ok=True)
        months_dir = folder / "по_месяцам"
        months_dir.mkdir(parents=True, exist_ok=True)

        # full year
        doc = build_document(meta, year, months)
        full_path = folder / f"plan_proverok_{safe_name(sro)}_2026.docx"
        doc.save(str(full_path))
        total = sum(len(v) for v in months.values())
        summary.append(f"{sro}: год={total} -> {full_path.name}")

        # by month
        for month in MONTH_ORDER:
            items = months.get(month) or []
            if not items:
                continue
            doc_m = build_document(meta, year, {month: items}, months_filter=[month])
            mpath = months_dir / f"{safe_name(sro)}_2026_{MONTH_FILE[month]}.docx"
            doc_m.save(str(mpath))
            summary.append(f"  {month}: {len(items)} -> {mpath.name}")

    readme = out_dir / "README.txt"
    readme.write_text(
        "Планы проверок собраны автоматически из Excel «ИТОГИ КОНТРОЛЯ 2026».\n"
        "СРО определено по цвету строки (легенда в 1-й строке).\n"
        "Оформление: Times New Roman 12, альбомный лист, без серой заливки — "
        "как официальные plan_proverok_*.doc.\n"
        "В каждом каталоге СРО:\n"
        "  - plan_proverok_<СРО>_2026.docx — весь год,\n"
        "  - папка по_месяцам — отдельный файл на каждый месяц.\n"
        "\n"
        "Важно: в Excel цветами размечены 9 СРО (ОГПС, ОГПП, ОСО, ОГПО,\n"
        "ГеоИндустрия, МГЕО, СПРОФ, ПРИИС, ОПП). Отдельные планы ОСОТ/НОСО/\n"
        "МОТС/ОСОВС/ОСОЕС/ГСП из zip здесь не выделяются — у них нет своего\n"
        "цвета в этой таблице (обычно ведутся отдельными выгрузками).\n"
        "\n"
        "Перед публикацией проверьте шапку (протокол) и состав строк.\n",
        encoding="utf-8",
    )

    report = out_dir / "_отчет.txt"
    report.write_text("\n".join(summary), encoding="utf-8")
    return summary


def main():
    desk = Path(r"C:\Users\User\OneDrive\Desktop")
    ap = argparse.ArgumentParser()
    ap.add_argument("--xlsx", default=str(desk / "ИТОГИ КОНТРОЛЯ 2026.xlsx"))
    ap.add_argument("--out", default=str(desk / "Планы_проверок_2026_Word"))
    ap.add_argument("--year", type=int, default=2026)
    ap.add_argument("--sro", default=None, help="Только одно СРО, напр. СПРОФ")
    args = ap.parse_args()
    summary = export(Path(args.xlsx), Path(args.out), year=args.year, only_sro=args.sro)
    print("DONE", args.out)
    print(f"files groups: {len([x for x in summary if not x.startswith('  ')])}")
    for line in summary[:40]:
        print(line)
    if len(summary) > 40:
        print("...")


if __name__ == "__main__":
    main()
