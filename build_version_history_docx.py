# -*- coding: utf-8 -*-
"""Сборка Word: история версий СРО-бота. Запуск: py build_version_history_docx.py"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "ISTORIYA_VERSIY_SRO_BOT.docx"
CURRENT = "2026-07-23-v56"


def _set_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)


def build() -> Path:
    doc = Document()
    _set_defaults(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Telegram-бот СРО Ассоциации «ГЕН»")
    tr.bold = True
    tr.font.size = Pt(16)
    tr.font.name = "Times New Roman"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("История версий: от v1 к v56")
    sr.bold = True
    sr.font.size = Pt(14)
    sr.font.name = "Times New Roman"

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        f"Актуальная версия: {CURRENT}\n"
        f"Дата документа: {date.today().strftime('%d.%m.%Y')}"
    )
    mr.font.size = Pt(11)
    mr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    mr.font.name = "Times New Roman"

    doc.add_paragraph()
    _para(
        doc,
        "Важно. Отдельного журнала на каждую из 56 версий не велось. "
        "Ниже — то, что удалось восстановить по файлам проекта и рабочим заметкам. "
        "Промежутки без записей помечены как «пробел».",
        italic=True,
    )

    # --- Эпохи ---
    _heading(doc, "1. Эпохи и вехи (v1 → v49)", 1)
    _para(
        doc,
        "Ранние номера сгруппированы в этапы: так честнее, чем выдумывать текст "
        "для каждой промежуточной версии.",
    )

    eras = [
        (
            "v1–v10 — каркас",
            [
                "Появление бота в Telegram: меню и базовые команды",
                "Первые FAQ и тексты под Ассоциацию «ГЕН»",
                "Локальный запуск одним процессом",
            ],
        ),
        (
            "v11–v21 — полезный продукт",
            [
                "Поиск организации по ИНН и названию",
                "Планы проверок из файлов Word (папка plany)",
                "FAQ: вступающим, членам, НОК/НРС, партнёры",
                "Контакты отделов (с паролем)",
                "ИИ-помощник со ссылками на официальные разделы сайта",
            ],
        ),
        (
            "v22 (19.07.2026) — презентация",
            [
                "Зафиксирован документ PREZENTACIYA_SRO_BOT.md",
                "Описаны меню, FAQ, ИИ, бланки, реестр",
                "Формула: сайт — источник истины, бот — быстрый вход",
            ],
        ),
        (
            "v23–v30 — пробел",
            ["Точечных записей по номерам версий не сохранилось"],
        ),
        (
            "v31 (20.07.2026) — IT-контур",
            [
                "Упоминание в IT_ZAYAVKA_DEPLOY.md",
                "Правило: номер в логе бота = номер, переданный IT",
            ],
        ),
        (
            "v32–v37 — пробел",
            ["Промежуточные правки без отдельного журнала"],
        ),
        (
            "v38 (20.07.2026) — прогон куратора",
            [
                "Упоминание в KURATOR_BOT_ROL.md и IT-документах",
                "Перед прогоном сверять версию в консоли",
            ],
        ),
        (
            "v39–v41 — пробел",
            ["Короткий промежуток до v42"],
        ),
        (
            "v42 (21.07.2026) — срез перед скачком",
            [
                "Последняя известная метка перед пакетом доработок к v50",
                "Далее — экосистема 15 СРО и бланки без подмены ОГПС",
            ],
        ),
        (
            "v43–v49 — экосистема 15 СРО (номера внутри — пробел)",
            [
                "15 комплектов бланков с официальных сайтов СРО",
                "Выбор СРО, если у организации несколько членств",
                "Убрана подмена «все строители → бланки ОГПС»",
                "Скрипты обновления бланков (discover / sync)",
                "Контекст: строители / проектировщики / изыскания",
                "Описание для руководства: DLYA_RUKOVODSTVA_BOTA.md",
            ],
        ),
    ]

    for title_text, bullets in eras:
        _heading(doc, title_text, 2)
        _bullets(doc, bullets)

    # --- Детальный журнал ---
    _heading(doc, "2. Детальный журнал (сохранился) — v50 → v56", 1)
    _para(
        doc,
        "С 22–23 июля 2026 версии фиксировались в коде (BOT_VERSION) "
        "и в рабочих заметках подробнее.",
    )

    details = [
        (
            "v50 — 22.07.2026 — база для отчётов",
            [
                "Стабильный срез для OTCHET_* и сборщиков Word",
                "Рабочий каркас: реестр, планы, бланки, ИИ, FAQ",
            ],
        ),
        (
            "v51 — 22.07.2026 — онбординг и фидбек",
            [
                "/start: ввод ИНН или «Пропустить (вступаю / без ИНН)»",
                "Кнопки «назад» к меню и к выбору СРО в бланках",
                "Ссылки планов/контроля с учётом выбранного СРО",
                "Кнопка «Ответ не помог» только под ответами ИИ",
                "Лог ожиданий пользователя → feedback_questions.jsonl",
                "Исправление ошибки при выборе СРО (UnboundLocalError)",
            ],
        ),
        (
            "v52 — 22.07.2026 — направление для вступающих",
            [
                "После «Пропустить» — выбор: строители / проектировщики / изыскания",
                "Контекст: ОГПС · ОГПП · ГеоИндустрия",
                "Без выбора направления главное меню не открывается",
                "Формулировки — на согласование с руководством",
            ],
        ),
        (
            "v53 — 23.07.2026 — пилот СПРОФ: информационный лист",
            [
                "Только СРО СПРОФ + бланк «Информационный лист»",
                "Автоподстановка регистрационного номера и ИНН из реестра",
                "Подставляется также наименование организации",
                "Без выбранной организации — пустой шаблон как раньше",
            ],
        ),
        (
            "v54 — 23.07.2026 — больше полей в инфолисте",
            [
                "Юридический, фактический и почтовый адрес (из «Местонахождение»)",
                "Руководитель организации",
                "Страховая компания и страховая сумма",
                "По-прежнему только пилот СПРОФ",
            ],
        ),
        (
            "v55 — 23.07.2026 — заявление о внесении изменений",
            [
                "Автозаполнение бланка заявления (СПРОФ)",
                "ИНН и ОГРН по клеткам",
                "Фактический адрес и руководитель (должность | ФИО)",
                "Наименования и адреса также подставляются",
            ],
        ),
        (
            "v56 — 23.07.2026 — прозрачность источника данных",
            [
                "В подписи к файлу — ссылка на реестр sprofproekt.ru/reestr/",
                "Ссылка на сайт СРО",
                "Явный текст: обязательно проверить сведения перед подачей",
                "Действует для информационного листа и заявления",
            ],
        ),
    ]

    for title_text, bullets in details:
        _heading(doc, title_text, 2)
        _bullets(doc, bullets)

    _heading(doc, "3. Как вести историю дальше", 1)
    _bullets(
        doc,
        [
            "При каждом заметном релизе поднимать BOT_VERSION в bot_FINAL_GOLD.py",
            "В тот же день дописывать 3–5 строк в VERSION_HISTORY.md",
            "При необходимости пересобрать этот Word: py build_version_history_docx.py",
        ],
    )

    foot = doc.add_paragraph()
    fr = foot.add_run(
        "\nДокумент подготовлен для внутренней демонстрации руководству. "
        "Источник: проект GOLD, файл VERSION_HISTORY.md."
    )
    fr.italic = True
    fr.font.size = Pt(10)
    fr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    fr.font.name = "Times New Roman"

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"OK: {path}")
