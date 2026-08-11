# -*- coding: utf-8 -*-
"""ТЗ на англоязычную версию бота — DOCX для руководства. python build_otchet_en_tz_docx.py"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml.ns import qn

OUT = Path(__file__).resolve().parent / "OTCHET_TZ_ANGLIYSKIY_BOT.docx"


def setup(doc: Document) -> None:
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(1.5)
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(12)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def main() -> None:
    doc = Document()
    setup(doc)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(
        "Технико-организационное предложение\n"
        "Англоязычная версия Telegram-бота СРО «ГЕН»"
    )
    r.bold = True
    r.font.size = Pt(14)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        f"Для согласования с руководством · {date.today().strftime('%d.%m.%Y')}\n"
        "Статус: предложение (не начато без решения руководства)"
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "Документ описывает зачем нужен английский канал, два варианта объёма (A/B), "
        "роли куратора и IT, риски и критерии успеха. Русскоязычный бот остаётся основным "
        "и юридически значимым; английский — навигация и ориентиры для не русскоязычных пользователей."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_page_break()

    doc.add_heading("1. Зачем английская версия", 1)
    for item in [
        "Иностранные партнёры, подрядчики, консультанты — быстрый контакт и понимание, что такое СРО «ГЕН».",
        "Компании из регионов с иностранным участником — первичный канал на английском.",
        "Имидж Ассоциации: цифровой сервис не только на русском.",
        "Не дублирует юридическую силу сайта: полные тексты норм и FAQ — на srogen.ru (RU).",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2. Два варианта объёма (выбор руководства)", 1)

    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    hdr = ["", "Вариант A — «мини-EN»", "Вариант B — «рабочий EN»", "Вариант C — полный паритет"]
    for i, h in enumerate(hdr):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
    rows = [
        ("Срок (ориентир)", "1–2 недели", "1–2 месяца", "много месяцев + юр. вычитка"),
        (
            "Содержание",
            "8–10 пунктов меню: About SRO, Contacts, Partners, How to join → ссылки; "
            "дисклеймер EN",
            "A + краткие EN-тексты: взносы, документы, NOK/NRS, проверки; "
            "ИНН-поиск без изменений",
            "Перевод 48 FAQ, всех блоков ИИ, справочника",
        ),
        ("Рекомендация", "Пилот", "После статистики A", "Только при явной потребности"),
    ]
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = val
    doc.add_paragraph()

    doc.add_paragraph(
        "Рекомендация куратора: начать с варианта A, через 2–3 месяца оценить запросы "
        "и решить, нужен ли B."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading("3. Два бота или один (технически)", 1)
    t2 = doc.add_table(rows=3, cols=3)
    t2.style = "Table Grid"
    for i, h in enumerate(["", "Два бота (@…_ru и @…_en)", "Один бот, выбор 🇷🇺/🇬🇧"]):
        t2.rows[0].cells[i].text = h
    t2.rows[1].cells[0].text = "Плюсы"
    t2.rows[1].cells[1].text = "Проще для пользователя; разные ссылки на сайте; проще аналитика"
    t2.rows[1].cells[2].text = "Один токен в BotFather; одна точка входа"
    t2.rows[2].cells[0].text = "Минусы"
    t2.rows[2].cells[1].text = "Два токена, два процесса на VPS (или один скрипт — два токена)"
    t2.rows[2].cells[2].text = "Сложнее UX; риск перепутать язык"
    doc.add_paragraph()
    doc.add_paragraph(
        "Рекомендация IT/куратору: два бота для пилота A (меньше путаницы). "
        "Регистрация второго бота — BotFather, токен хранит IT на сервере."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading("4. Обязательный дисклеймер (английский текст для бота)", 1)
    box = doc.add_paragraph()
    box.paragraph_format.left_indent = Cm(1)
    box.add_run(
        "This bot provides general guidance and links to the official website of "
        "SRO Association «GEN» (srogen.ru). Legal documents and binding information "
        "are published in Russian on the official website. For individual cases, "
        "please contact the Association: +7 (495) 775-81-11, info@srogen.ru."
    ).italic = True

    doc.add_heading("5. Роли", 1)
    t3 = doc.add_table(rows=5, cols=2)
    t3.style = "Table Grid"
    t3.rows[0].cells[0].text = "Кто"
    t3.rows[0].cells[1].text = "Задачи"
    for p in t3.rows[0].cells[0].paragraphs + t3.rows[0].cells[1].paragraphs:
        for run in p.runs:
            run.bold = True
    roles = [
        ("Руководство", "Утвердить вариант A или B; нужен ли EN вообще; бюджет перевода"),
        (
            "Куратор",
            "ТЗ на тексты EN; согласование с аппаратом; приёмка в Telegram; "
            "не программирование",
        ),
        (
            "IT / подрядчик",
            "Второй бот на VPS; systemd; секреты; деплой по IT_ZAYAVKA_DEPLOY.md",
        ),
        (
            "Переводчик + аппарат",
            "Вариант B: вычитка EN-формулировок (не только машинный перевод)",
        ),
    ]
    for i, (who, task) in enumerate(roles, 1):
        t3.rows[i].cells[0].text = who
        t3.rows[i].cells[1].text = task
    doc.add_paragraph()

    doc.add_heading("6. Что не переводится / не меняется", 1)
    for item in [
        "Поиск по ИНН, реестр, планы проверок — логика та же; подписи полей можно EN.",
        "Ссылки на srogen.ru — без изменения URL (страницы RU).",
        "Телефонный справочник — ФИО на русском; EN-инструкция «Russian staff directory».",
        "Партнёры — названия СРО как на сайте.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("7. Риски", 1)
    for item in [
        "Автоперевод норм — ошибки → репутационный риск (только согласованные тексты).",
        "Пользователь считает EN юридически полным → дисклеймер в каждом блоке.",
        "Два языка — двойная регрессия (чек-листы куратора).",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("8. Критерии успеха пилота (вариант A)", 1)
    for item in [
        "Бот EN отвечает 24/7 на VPS (как RU).",
        "Меню EN покрывает: About, Join (link), Contacts, Partners, FAQ (link), disclaimer.",
        "Нет жалоб на «бот соврал по закону» — только ссылки + контакты.",
        "За 3 месяца: зафиксировано ≥ N обращений / или решение не развивать EN.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("9. Порядок запуска (если «да»)", 1)
    steps = [
        "Решение руководства: вариант A или B.",
        "BotFather: создать EN-бота, токен → IT.",
        "Куратор: тексты EN (1–2 страницы Word), согласование.",
        "IT: деплой (форк texts_en + тот же код или env LANG=en).",
        "Куратор: прогон сценариев; ссылка t.me/… на сайт при необходимости.",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}", style="List Number")

    doc.add_heading("10. Связь с текущим проектом", 1)
    doc.add_paragraph(
        "Русскоязычный бот (версия 2026-07-22-v50) — основной продукт. "
        "Английский — добавление после стабильного VPS RU. "
        "См. также: OTCHET_KURATOR_I_BOT_SRO_GEN.docx, OTCHET_ZACHEM_BOT_REGIONY.docx."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.add_run(
        "Подготовлено для внутреннего согласования · Куратор бота СРО «ГЕН»"
    ).font.size = Pt(10)

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
