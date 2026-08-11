#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Документы к завтра: архитектура ПДн (Москва + OpenRouter) и чеклист переезда."""
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

DESKTOP = Path(r"C:\Users\User\Desktop")


def new_doc() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)
    return doc


def font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def title_block(doc, main, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(main)
    font(r, 16, bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    font(r2, 13, bold=True)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
    font(r3, 11, italic=True)


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        font(run, 14 if level == 1 else 13, bold=True)


def para(doc, text, *, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    font(run, italic=italic)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        font(run)
        p.paragraph_format.space_after = Pt(3)


def table2(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, htext in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = htext
        for p in cell.paragraphs:
            for run in p.runs:
                font(run, bold=True, size=11)
    for row_data in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = val
            for p in cells[i].paragraphs:
                for run in p.runs:
                    font(run, size=11)
    doc.add_paragraph()


def write_architecture() -> Path:
    doc = new_doc()
    title_block(
        doc,
        "ГДЕ ЧТО ЛЕЖИТ: БОТ СРО И ИИ",
        "Ответы на вопросы (в т.ч. юридические)\n"
        "схема как у Minecraft: тело в Москве, ИИ-сервис снаружи",
    )
    para(
        doc,
        "Документ для внутренних ответов руководству и юристу. "
        "Не заменяет юридическое заключение.",
        italic=True,
    )

    h(doc, "1. Короткий ответ «на завтра»")
    para(
        doc,
        "Тело бота и вся информация о пользователях и реестре будут храниться "
        "на сервере в Москве (Российская Федерация). "
        "Ключ и сервис OpenRouter (ИИ) находятся за рубежом — это отдельный "
        "облачный сервис, без которого ответы ИИ не работают. "
        "Ключ можно (и нужно) хранить в конфиге на московском сервере; "
        "запросы при этом всё равно уходят на инфраструктуру OpenRouter за границей — "
        "так устроен сам сервис, «переехать OpenRouter в Москву» нельзя.",
    )
    bullets(
        doc,
        [
            "ПДн и рабочие данные бота → Москва (РФ).",
            "OpenRouter → зарубежный API (как у проекта Minecraft).",
            "Telegram → мессенджер; бот отвечает из Москвы.",
        ],
    )

    h(doc, "2. Как сейчас и как будет")
    table2(
        doc,
        ["Что", "Сейчас", "Целевая схема"],
        [
            (
                "Сервер бота (код, systemd)",
                "VPS Timeweb, Амстердам (NL), IP 147.45.225.70",
                "VPS в Москве (РФ) — отдельный или рядом с MC-сервером",
            ),
            (
                "Журнал пользователей bot_users.json",
                "На NL-сервере",
                "Только на сервере в Москве",
            ),
            (
                "Кэш реестра, контекст СРО, логи",
                "На NL-сервере",
                "На сервере в Москве",
            ),
            (
                "Ключ OPENROUTER_API_KEY в config_keys.py",
                "Файл на NL-сервере; API OpenRouter за рубежом",
                "Файл на Москве; API OpenRouter по-прежнему за рубежом",
            ),
            (
                "Minecraft (для сравнения)",
                "Игра / сервер — Москва (201.24.125.236); ИИ-мост — NL",
                "Та же логика: тяжёлые/личные данные в РФ, ИИ-вызов снаружи",
            ),
        ],
    )

    h(doc, "3. Важная уточняющая фраза (чтобы не путать)")
    para(
        doc,
        "«Ключ лежит в Нидерландах» и «сервис OpenRouter в Нидерландах» — разные вещи.",
    )
    bullets(
        doc,
        [
            "Ключ — это строка в нашем файле config_keys.py. Её мы кладём на московский сервер "
            "(как пароль в сейфе в Москве).",
            "Сервис OpenRouter — чужие компьютеры за границей. Когда бот отвечает через ИИ, "
            "он на секунду отправляет туда текст вопроса и получает ответ.",
            "Если «перенести OpenRouter в Москву» — такого продукта у нас нет: "
            "без их API ответы ИИ в боте работать не будут. Запасной вариант — Groq "
            "(тоже облако за рубежом) или отключить ИИ и оставить только меню/реестр.",
        ],
    )

    h(doc, "4. Что сказать юристу / по 152-ФЗ")
    bullets(
        doc,
        [
            "Оператор (Ассоциация) обрабатывает ПДн через бот: Telegram ID, имя, username, "
            "служебный контекст (файл на сервере в РФ после переезда).",
            "Первичная база и файлы бота — на территории РФ (Москва) — для локализации это "
            "правильная целевая картина.",
            "Отдельно остаются: (а) сам Telegram как иностранная платформа; "
            "(б) при использовании ИИ — краткая передача текста запроса на OpenRouter. "
            "Это оформляется политикой, согласием и при необходимости — "
            "оценкой трансграничной передачи у юриста.",
            "Рекламных рассылок и приёма платежей в боте нет.",
            "Нужны: политика ПДн, текст согласия в /start, вопрос про уведомление Роскомнадзора.",
        ],
    )

    h(doc, "5. Аналогия с Minecraft (удобно объяснять)")
    para(
        doc,
        "У Minecraft уже так: игровой сервер и основные данные — в Москве; "
        "ИИ-мост и часть сервисов — на VPS в NL, потому что так удобнее к OpenRouter. "
        "Для бота СРО мы делаем ещё жёстче по смыслу ПДн: "
        "сам бот и журнал пользователей переезжают в Москву, "
        "а наружу уходит только вызов ИИ по необходимости.",
    )

    h(doc, "6. Что не является проблемой")
    bullets(
        doc,
        [
            "То, что компания Timeweb российская, а старый VPS был в NL — бывает; "
            "важно место сервера с данными, не бренд хостера.",
            "Хранение ключа OpenRouter в Москве не ломает бота — "
            "ломает только отсутствие доступа к их API или неверный ключ.",
        ],
    )

    h(doc, "7. Статус переезда")
    para(
        doc,
        "Подготовлен чеклист переноса (отдельный файл). "
        "Фактический переезд — на VPS в Москве (можно использовать имеющийся "
        "московский сервер 201.24.125.236 отдельным каталогом /opt/sro-bot "
        "либо новый VPS Timeweb «Москва»). "
        "До переезда production всё ещё на 147.45.225.70 (NL).",
        italic=True,
    )

    out = DESKTOP / "Бот_СРО_где_лежат_данные_Москва_и_OpenRouter.docx"
    doc.save(out)
    return out


def write_migrate_checklist() -> Path:
    doc = new_doc()
    title_block(
        doc,
        "ЧЕКЛИСТ: ПЕРЕЕЗД БОТА СРО В МОСКВУ",
        "Чтобы тело и данные были в РФ\n"
        "(OpenRouter остаётся внешним API)",
    )

    h(doc, "1. Цель")
    para(
        doc,
        "Перенести /opt/sro-bot с NL (147.45.225.70) на сервер в Москве. "
        "После переезда: пользователи, реестр, логи — в РФ; "
        "ИИ по-прежнему ходит в OpenRouter.",
    )

    h(doc, "2. Варианты сервера в Москве")
    bullets(
        doc,
        [
            "Вариант A (быстрее): тот же VPS, что Minecraft — 201.24.125.236, "
            "отдельная папка /opt/sro-bot и отдельный systemd (не мешать MC).",
            "Вариант B: новый VPS Timeweb с локацией Москва — чище для Ассоциации.",
        ],
    )
    para(
        doc,
        "Рекомендация: для «юридически красиво» лучше отдельный VPS «Москва» "
        "или явно зафиксировать, что MC и бот СРО — разные каталоги на одном РФ-сервере.",
        italic=True,
    )

    h(doc, "3. Что переносим")
    bullets(
        doc,
        [
            "Код: bot_FINAL_GOLD.py и модули (.py), vps/, blanki/plany или sro_data.",
            "Секреты: config_keys.py (ТОЛЬКО руками, не из Windows-пути вслепую) — "
            "SRO_FILES_DIR на сервере должен быть /opt/sro-bot/sro_data или как сейчас на VPS.",
            "Данные: reestr_cache.json, bot_users.json, user_sro_context.json, "
            "nrs_link_mode.json при наличии.",
            "Сервисы: sro-bot.service, таймеры reestr-sync и site-health, "
            "скрипты + maintenance_stub.",
            "Не копировать venv с NL — создать новый venv на Москве.",
        ],
    )

    h(doc, "4. Порядок работ (чтобы не поймать 409)")
    bullets(
        doc,
        [
            "1) Поднять Москву: каталог, venv, зависимости, config_keys, файлы данных.",
            "2) Прогнать py_compile и короткий smoke (импорт модулей).",
            "3) Поставить systemd, НО не запускать бота ещё.",
            "4) На NL: systemctl stop sro-bot (и stub, если есть).",
            "5) На Москве: systemctl start sro-bot → проверить active и ответ в Telegram.",
            "6) Включить таймеры sync и health на Москве; на NL — disable таймеры бота.",
            "7) Обновить шпаргалки деплоя (IP Москвы вместо 147.45.225.70).",
            "8) NL: оставить только то, что нужно MC (ai-bridge и т.д.), бот СРО не запускать.",
        ],
    )

    h(doc, "5. Проверка «для завтра»")
    bullets(
        doc,
        [
            "curl/ipinfo по IP бота → country RU / город Москва (или РФ-регион).",
            "/start в Telegram отвечает.",
            "Поиск организации по ИНН работает.",
            "ИИ-помощник отвечает (значит ключ OpenRouter с московского сервера достучался).",
            "Файлы bot_users.json и reestr_cache.json лежат на Москве.",
        ],
    )

    h(doc, "6. Формулировка для руководства")
    para(
        doc,
        "«Сервер бота и все сохранённые данные — в Москве. "
        "Для умных ответов бот обращается к облачному сервису OpenRouter "
        "(как и в других наших проектах). "
        "Сам OpenRouter в Россию не переносится — это их сервис; "
        "без него блок ИИ не работает. Меню, реестр и бланки работают и без ИИ.»",
    )

    h(doc, "7. Риски при переезде")
    bullets(
        doc,
        [
            "Два бота с одним токеном → ошибка 409. Поэтому NL гасим до старта Москвы.",
            "Неверный SRO_FILES_DIR → не найдутся бланки.",
            "Забыли таймеры → реестр не обновится ночью.",
            "Мало RAM на общем MC-сервере → лучше отдельный VPS.",
        ],
    )

    out = DESKTOP / "Чеклист_переезд_бота_СРО_в_Москву.docx"
    doc.save(out)
    return out


def patch_legal_report() -> Path | None:
    """Пересобрать отчёт ПДн с честной географией."""
    # Re-use simplified correction file
    doc = new_doc()
    title_block(
        doc,
        "ОТЧЁТ (уточнение)",
        "География сервера бота СРО и персональные данные",
    )
    para(
        doc,
        "Ранее в черновике ошибочно указывалось, что production-VPS находится в РФ. "
        "Фактически на дату проверки IP 147.45.225.70 (Timeweb) определяется как "
        "Амстердам, Нидерланды.",
    )
    para(
        doc,
        "Целевое состояние: перенос бота и хранилища данных в Москву (РФ); "
        "OpenRouter остаётся внешним API. Подробности — в файлах "
        "«Бот_СРО_где_лежат_данные_Москва_и_OpenRouter.docx» и "
        "«Чеклист_переезд_бота_СРО_в_Москву.docx».",
    )
    para(
        doc,
        "Полный отчёт по 152-ФЗ / 38-ФЗ / 54-ФЗ см. "
        "«Отчёт_бот_СРО_ПДн_и_законность.docx» — читать вместе с этим уточнением.",
        italic=True,
    )
    out = DESKTOP / "Уточнение_сервер_бота_не_РФ_а_NL_цель_Москва.docx"
    doc.save(out)
    return out


def main() -> None:
    a = write_architecture()
    b = write_migrate_checklist()
    c = patch_legal_report()
    for path in (a, b, c):
        print(path)
        print(path.stat().st_size)


if __name__ == "__main__":
    main()
